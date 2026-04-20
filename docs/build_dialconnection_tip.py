#!/usr/bin/env python3
"""
DialConnection LLC - Technical Implementation Plan
ThriveCloud Migration from TierPoint — CON-094865 / Q-69941
Rebuilt April 20, 2026 — Astra-format TIP with full Excel inventory
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CLIENT_NAME = 'DialConnection LLC'
REPORT_DATE = 'April 20, 2026'
REPORT_SUBTITLE = 'ThriveCloud Migration — Technical Implementation Plan'
OUTPUT_PATH = '/Users/josheddolls/Documents/Reports/DialConnection LLC - Technical Implementation Plan.docx'
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'thrive-logo.png')


# ---------------------------------------------------------------------------
# Core style helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color='000000', sz='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_padding(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        item = OxmlElement(f'w:{side}')
        item.set(qn('w:w'), str(val))
        item.set(qn('w:type'), 'dxa')
        mar.append(item)
    tcPr.append(mar)


def set_para_spacing(para, before=0, after=60, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(p, before=60, after=60)


def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)


def add_doc_header(doc, date, client, subtitle):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(LOGO_PATH, width=Inches(1.8))
    set_para_spacing(p, before=0, after=120)

    p = doc.add_paragraph()
    r = p.add_run(date)
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    set_para_spacing(p, before=0, after=40)

    p = doc.add_paragraph()
    r = p.add_run(client)
    r.font.name = 'Calibri'
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    set_para_spacing(p, before=0, after=60)

    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    set_para_spacing(p, before=0, after=40)

    p = doc.add_paragraph()
    r = p.add_run('Prepared by Thrive')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    set_para_spacing(p, before=0, after=100)

    add_hrule(doc)


def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.bold = True
    r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    if level == 1:
        r.font.size = Pt(14)
        set_para_spacing(p, before=240, after=80)
    elif level == 2:
        r.font.size = Pt(12)
        set_para_spacing(p, before=180, after=60)
    else:
        r.font.size = Pt(11)
        set_para_spacing(p, before=140, after=40)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    set_para_spacing(p, before=0, after=80)
    return p


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    set_para_spacing(p, before=0, after=40)
    return p


def add_checklist_item(doc, text, completed=False):
    p = doc.add_paragraph(style='List Bullet')
    indicator = '[X] ' if completed else '[ ] '
    r1 = p.add_run(indicator)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r1.bold = True
    if completed:
        r1.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
    else:
        r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    set_para_spacing(p, before=0, after=40)
    return p


def add_sub_note(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    pPr.append(ind)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(9.5)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_para_spacing(p, before=0, after=60)
    return p


def add_table(doc, headers, rows, col_widths, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    hdr.height = Inches(0.33)
    for index, header in enumerate(headers):
        cell = hdr.cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, '007BFF')
        set_cell_border(cell, color='007BFF', sz='4')
        set_cell_padding(cell, top=60, bottom=60, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, before=0, after=0)
        r = p.add_run(header)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index + 1]
        row.height = Inches(0.27)
        bg = 'F7F7F7' if row_index % 2 == 0 else 'FFFFFF'
        for col_index, value in enumerate(row_data):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            set_cell_border(cell, color='CCCCCC', sz='4')
            set_cell_padding(cell, top=40, bottom=40, left=100, right=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, before=0, after=0)
            r = p.add_run(str(value) if value is not None else '')
            r.font.name = 'Calibri'
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    for row in table.rows:
        for col_index, width in enumerate(col_widths):
            row.cells[col_index].width = Inches(width)

    doc.add_paragraph()
    return table


def add_footer(doc):
    add_hrule(doc)
    p = doc.add_paragraph()
    r = p.add_run('Thrive  ·  Managed Services  ·  support@thrivenextgen.com')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=80, after=0)


# ---------------------------------------------------------------------------
# Server inventory data (from April 2026 workbook)
# ---------------------------------------------------------------------------

CHICAGO_SERVERS = [
    # Hostname, Role, OS, vCPU, RAM, Disk, DRaaS, Migrate, Notes
    ('AL-IDS-Appliance-CHH', 'Platform', 'Linux', '4', '4', '58', 'No', 'N', 'Alert Logic IDS — replaced by MDR'),
    ('AL-LOG-Appliance-CHH', 'Platform', 'Linux', '2', '2', '56', 'No', 'N', 'Alert Logic LOG — replaced by MDR'),
    ('DCloud3CApp002A', 'Production', 'Windows 2022', '8', '12', '192', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp004', 'Production', 'Windows 2022', '2', '8', '68', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp009U', 'Production', 'Windows 2022', '4', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp017U', 'Production', 'Windows 2022', '4', '12', '174', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp025', 'Production', 'Windows 2022', '1', '4', '154', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp026U', 'Production', 'Windows 2022', '4', '12', '162', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp028A', 'Production', 'Windows 2022', '12', '12', '192', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp029U', 'Production', 'Windows 2022', '2', '8', '68', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp031U', 'Production', 'Windows 2022', '4', '12', '132', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp043', 'Production', 'Windows 2022', '8', '10', '174', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp044', 'Production', 'Windows 2022', '2', '8', '128', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CApp046', 'Production', 'Windows 2022', '4', '10', '130', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CGPX008', 'Production', 'Windows 2022', '4', '8', '128', 'Yes', 'Y', 'App Server'),
    ('DCloud3CGPX009', 'Production', 'Windows 2022', '4', '8', '128', 'Yes', 'Y', 'App Server'),
    ('DCloud3CGPX012', 'Production', 'Windows 2022', '2', '8', '166', 'Yes', 'Y', 'App Server'),
    ('DCloud3CGPX021U', 'Production', 'Windows 2022', '2', '8', '136', 'Yes', 'Y', 'App Server'),
    ('DCloud3CGPX026', 'Production', 'Windows 2022', '2', '8', '136', 'Yes', 'Y', 'App Server'),
    ('DCloud3CPDC', 'Production', 'Windows 2022', '4', '8', '168', 'Yes', 'Y', 'Primary Domain Controller'),
    ('DCloud3CRDSCB1', 'Production', 'Windows 2022', '2', '8', '68', 'Yes', 'Y', 'RDS Connection Broker — Public Farm'),
    ('DCloud3CRDSDC1', 'Production', 'Windows 2022', '6', '32', '432', 'Yes', 'Y', 'Standalone RDS — Internal Only'),
    ('DCloud3CRDSGW1', 'Production', 'Windows 2022', '2', '8', '68', 'Yes', 'Y', 'RDS Gateway — Public Farm'),
    ('DCLOUD3CRDSSH1', 'Production', 'Windows 2022', '4', '8', '68', 'Yes', 'Y', 'RDS Session Host — Public Farm'),
    ('DCLOUD3CRDSSH2', 'Production', 'Windows 2022', '4', '8', '68', 'Yes', 'Y', 'RDS Session Host — Public Farm'),
    ('DCloud3CSDC1', 'Production', 'Windows 2022', '2', '8', '3,058', 'Yes', 'Y', 'SDC and File Server'),
    ('DCloud3CSQLCST1', 'Production', 'Windows 2022', '24', '192', '6,915', 'Yes', 'Y', 'SQL Server — Primary (24vCPU/192GB)'),
    ('DCloud3CSQLCST2', 'Production', 'Windows 2022', '20', '128', '1,799', 'Yes', 'Y', 'SQL Server'),
    ('DCloud3CSQLCST3', 'Production', 'Windows 2022', '16', '96', '2,057', 'Yes', 'Y', 'SQL Server'),
    ('DCloud3CWeb2', 'Production', 'Windows 2022', '8', '32', '288', 'Yes', 'Y', 'Web Server'),
    ('DCloud3CSQLDev1', 'Development', 'Windows 2022', '2', '8', '758', 'No', 'Y', '→ ATLANTA (dev consolidation)'),
    ('GatewayCHI001', 'Production', 'Linux Debian 12', '14', '32', '500', 'No', 'Y', 'SIP Gateway'),
    ('GatewayCHI002', 'Production', 'Linux Debian 12', '14', '32', '500', 'No', 'Y', 'SIP Gateway'),
    ('GatewayCHI003', 'Production', 'Linux Debian 12', '16', '32', '500', 'No', 'Y', 'SIP Gateway'),
    ('GatewayCHI004', 'Production', 'Linux Debian 12', '16', '32', '500', 'No', 'Y', 'SIP Gateway'),
    ('GatewayCHIDR001', 'DR', 'Linux Debian 12', '12', '32', '500', 'No', 'Y', 'SIP Gateway DR'),
    ('GatewayCHIDR002', 'DR', 'Linux Debian 12', '12', '32', '500', 'No', 'Y', 'SIP Gateway DR (review removal)'),
    ('Fortigate Firewall 01', 'Platform', 'Physical', '—', '—', '—', 'No', 'N', 'FortiGate 100F HA — replaced by 200F'),
    ('Fortigate Firewall 02', 'Platform', 'Physical', '—', '—', '—', 'No', 'N', 'FortiGate 100F HA — replaced by 200F'),
]

DALLAS_SERVERS = [
    ('AL-LOG-Appliance-DA2', 'Platform', 'Linux', '2', '2', '56', 'No', 'N', 'Alert Logic LOG — replaced by MDR'),
    ('AL-IDS-Appliance-DA2', 'Platform', 'Linux', '4', '4', '58', 'No', 'N', 'Alert Logic IDS — replaced by MDR'),
    ('DCloud3DApp001U', 'Production', 'Windows 2022', '6', '12', '162', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp007U', 'Production', 'Windows 2022', '2', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp023U', 'Production', 'Windows 2022', '2', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp024', 'Production', 'Windows 2022', '1', '4', '154', 'Yes', 'Y', 'Web Server (PST)'),
    ('DCloud3DApp034U', 'Production', 'Windows 2022', '12', '12', '212', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp035U', 'Production', 'Windows 2022', '4', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp037', 'Production', 'Windows 2022', '2', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp039U', 'Production', 'Windows 2022', '2', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp041', 'Production', 'Windows 2022', '2', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DApp042', 'Production', 'Windows 2022', '4', '8', '158', 'Yes', 'Y', 'Web Server'),
    ('DCloud3DGPX039U', 'Production', 'Windows 2022', '2', '8', '158', 'Yes', 'Y', 'App Server'),
    ('DCloud3DPDC', 'Production', 'Windows 2022', '4', '10', '110', 'Yes', 'Y', 'Primary Domain Controller'),
    ('DCloud3DRDSDC1', 'Production', 'Windows 2022', '6', '32', '288', 'Yes', 'Y', 'Standalone RDS — Internal Only'),
    ('DCLOUD3DRDSSH1', 'Production', 'Windows 2022', '2', '8', '68', 'Yes', 'Y', 'RDS Session Host — Public Farm'),
    ('DCLOUD3DRDSSH2', 'Production', 'Windows 2022', '2', '8', '68', 'Yes', 'Y', 'RDS Session Host — Public Farm'),
    ('DCloud3DRDSCB1', 'Production', 'Windows 2022', '2', '4', '124', 'Yes', 'Y', 'RDS Connection Broker — Public Farm'),
    ('DCloud3DRDSGW1', 'Production', 'Windows 2022', '2', '4', '64', 'Yes', 'Y', 'RDS Gateway — Public Farm'),
    ('DCloud3DSDC1', 'Production', 'Windows 2022', '2', '8', '2,580', 'Yes', 'Y', 'SDC and File Server'),
    ('DCloud3DSQLEST1', 'Production', 'Windows 2022', '24', '192', '4,055', 'Yes', 'Y', 'SQL Server — EST Primary'),
    ('DCloud3DSQLEST2', 'Production', 'Windows 2022', '12', '96', '1,431', 'Yes', 'Y', 'SQL Server — EST Secondary'),
    ('DCloud3DSQLPST1', 'Production', 'Windows 2022', '8', '32', '3,383', 'Yes', 'Y', 'SQL Server — PST'),
    ('DCloud3DWeb2', 'Production', 'Windows 2022', '8', '32', '288', 'Yes', 'Y', 'Web Server'),
    ('GatewayDAL001', 'Production', 'Linux Debian 12', '12', '32', '500', 'Yes', 'Y', 'SIP Gateway'),
    ('GatewayDAL002', 'Production', 'Linux Debian 12', '12', '32', '500', 'Yes', 'Y', 'SIP Gateway'),
    ('GatewayDALDR001', 'DR', 'Linux Debian 12', '14', '32', '500', 'Yes', 'Y', 'SIP Gateway DR'),
    ('GatewayDALDR002', 'DR', 'Linux Debian 12', '14', '32', '500', 'Yes', 'Y', 'SIP Gateway DR (review removal)'),
    ('GatewayDALDR003', 'DR', 'Linux Debian 12', '16', '32', '500', 'Yes', 'Y', 'SIP Gateway DR (review removal)'),
    ('GatewayDALDR004', 'DR', 'Linux Debian 12', '16', '32', '500', 'Yes', 'Y', 'SIP Gateway DR (review removal)'),
    ('Fortigate Firewall 03', 'Platform', 'Physical', '—', '—', '—', 'No', 'N', 'FortiGate 100F HA — replaced by 200F'),
    ('Fortigate Firewall 04', 'Platform', 'Physical', '—', '—', '—', 'No', 'N', 'FortiGate 100F HA — replaced by 200F'),
    ('DC3DAppTemplate', 'Template', 'Linux Debian 12', '—', '—', '143', 'No', 'Y', 'Powered-off deployment template'),
    ('DC3DGPXTemplate', 'Template', 'Linux Debian 12', '—', '—', '143', 'No', 'Y', 'Powered-off deployment template'),
]

DALLAS_DEV_TO_ATL = [
    # Dev servers consolidated to Atlanta
    ('DC3DevPDC', 'Development', 'Windows 2022', '2', '4', '64', 'No', 'Y', 'Dev DC → Atlanta'),
    ('DC3DevDB1', 'Development', 'Windows 2022', '2', '16', '1,166', 'No', 'Y', 'Dev Database → Atlanta'),
    ('DC3DevGW001', 'Development', 'Linux Debian 12', '1', '2', '94', 'No', 'Y', 'Dev Gateway → Atlanta'),
    ('DC3DevGW002', 'Development', 'Linux Debian 12', '1', '2', '94', 'No', 'Y', 'Dev Gateway → Atlanta'),
    ('DC3DevGW003', 'Development', 'Linux Debian 12', '1', '2', '94', 'No', 'Y', 'Dev Gateway → Atlanta'),
    ('DC3DevGW004', 'Development', 'Linux Debian 12', '1', '2', '34', 'No', 'Y', 'Dev Gateway → Atlanta'),
    ('DC3DevQA001', 'Development', 'Windows 2022', '2', '8', '108', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevQA002', 'Development', 'Windows 2022', '4', '12', '112', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevQA003', 'Development', 'Windows 2022', '2', '8', '108', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevQA004', 'Development', 'Windows 2022', '4', '12', '112', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevQA005', 'Development', 'Windows 2022', '2', '8', '108', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevQA006', 'Development', 'Windows 2022', '4', '12', '112', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevQA007', 'Development', 'Windows 2022', '2', '4', '—', 'No', 'Y', 'QA Environment → Atlanta'),
    ('DC3DevWeb1', 'Development', 'Windows 2016', '6', '16', '644', 'No', 'Y', 'Dev Web Server → Atlanta'),
    ('DC3DevWeb2', 'Development', 'Windows 2022', '6', '16', '336', 'No', 'Y', 'Dev Web Server → Atlanta'),
    ('DC3DAppStaging', 'Development', 'Windows 2022', '2', '4', '64', 'No', 'Y', 'App Staging → Atlanta'),
    ('DC3DGPXStaging', 'Development', 'Windows 2022', '2', '4', '132', 'No', 'Y', 'GPX Staging → Atlanta'),
    ('DC3VBVM01', 'Development', 'Windows XP', '1', '4', '116', 'No', 'Y', 'Legacy Dev VM → Atlanta (confirm)'),
    ('DC3VBVM02', 'Development', 'Windows XP', '1', '4', '116', 'No', 'Y', 'Legacy Dev VM → Atlanta (confirm)'),
    ('DC3VBVM03', 'Development', 'Windows XP', '1', '4', '116', 'No', 'Y', 'VERIFY existence → Atlanta'),
    ('DC3VBVM04', 'Development', 'Windows XP', '1', '4', '116', 'No', 'Y', 'VERIFY existence → Atlanta'),
]

NASHVILLE_ATL_SERVERS = [
    ('DCloud3NPBX01', 'Prod — PBX', 'Linux', '16', '16', '400', 'Yes', 'Y', 'Shared PBX — CST Clients'),
    ('DCloud3NDRSQL1', 'DR', 'Windows 2022', '4', '6', '186', 'Yes', 'Y', 'DR SQL — CST PBX'),
    ('DCloud3N-Nagios', 'Monitoring', 'Linux', '1', '2', '66', 'Yes', 'Y', 'Nagios — will be replaced by Thrive monitoring'),
    ('NSH-DRGateway-PBX', 'DR', 'Linux', '1', '2', '34', 'Yes', 'Y', 'DR Gateway PBX — CST'),
    ('NSH Multi-Tenant Firewall', 'Platform', 'Physical', '—', '—', '—', 'No', 'N', 'FortiGate 80F HA — replaced by VDC virtual FW'),
]

LENEXA_BOS_SERVERS = [
    ('DCloud3KPBX01', 'Prod — PBX', 'Linux', '12', '12', '492', 'Yes', 'Y', 'Shared PBX — EST/PST Clients'),
    ('DCloud3LDRSQL1', 'DR', 'Windows 2022', '2', '4', '260', 'Yes', 'Y', 'DR SQL — EST/PST PBX'),
    ('LEN-DRGateway-PBX', 'DR', 'Linux', '1', '2', '36', 'Yes', 'Y', 'DR Gateway PBX — EST/PST'),
    ('LEN Multi-Tenant Firewall', 'Platform', 'Physical', '—', '—', '—', 'No', 'N', 'FortiGate 80F HA — replaced by VDC virtual FW'),
]

SERVER_HEADERS = ['Hostname', 'Role', 'OS', 'vCPU', 'RAM', 'Disk(GB)', 'DRaaS', 'Migrate', 'Notes']
SERVER_WIDTHS  = [1.45, 0.75, 1.2, 0.4, 0.4, 0.55, 0.5, 0.5, 0.85]  # = 6.6"


# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------

def build_tip():
    doc = Document()
    set_doc_margins(doc)
    add_doc_header(doc, REPORT_DATE, CLIENT_NAME, REPORT_SUBTITLE)

    # -----------------------------------------------------------------------
    # Executive Summary
    # -----------------------------------------------------------------------
    add_section_heading(doc, 'Executive Summary', 1)
    add_body(doc,
        'DialConnection LLC is a hosted PBX and VoIP service provider migrating its entire infrastructure '
        'from TierPoint (4 source datacenter locations) to Thrive ThriveCloud (4 target Virtual Datacenters '
        'in Chicago, Dallas, Atlanta, and Boston). This Technical Implementation Plan (TIP) defines the '
        'architecture, migration strategy, phased execution plan, and client/Thrive responsibilities '
        'for Service Order CON-094865, signed March 27, 2026.'
    )
    add_body(doc,
        'The migration encompasses 95 virtual machines and 4 physical firewall platforms across four sites. '
        'Chicago and Dallas are deployed on Thrive Dedicated Hosts with FortiGate 200F HA pairs. Atlanta '
        'and Boston are managed shared VDC environments. All production servers are backed up and replicated '
        'for DRaaS. The dev environment — currently split between Dallas and Chicago — is consolidated to '
        'Thrive Atlanta as a deliberate architectural decision. Migration is executed via Carbonite Migrate '
        'block-level replication under a 90-license contract (subject to Change Log reconciliation '
        'against the confirmed 95-VM inventory).'
    )
    add_bullet(doc, 'Client: ', 'DialConnection LLC — Keith Larson, President and COO')
    add_bullet(doc, 'Service Order: ', 'CON-094865 / Q-69941 — March 27, 2026')
    add_bullet(doc, 'Contract Term: ', '63 months (3 free months credit against 60-month base term)')
    add_bullet(doc, 'Net MRC: ', '$69,055.20/month after discount')
    add_bullet(doc, 'One-Time Net: ', '$62,685.74 implementation + $40,500.00 Carbonite licenses')
    add_bullet(doc, 'One-Time Credit: ', '$207,165.60 (~3 months MRC) applied at billing activation')
    add_bullet(doc, 'Thrive Project Manager: ', 'To be assigned at project kickoff')
    add_bullet(doc, 'Thrive Technical Lead: ', 'To be assigned at project kickoff')

    # -----------------------------------------------------------------------
    # Revision History
    # -----------------------------------------------------------------------
    add_section_heading(doc, 'Revision History', 1)
    add_table(doc,
        ['Rev', 'Author', 'Description', 'Date'],
        [
            ['1.0', 'Joshua Eddolls', 'Initial TIP — service order and architecture sections', 'April 7, 2026'],
            ['2.0', 'Joshua Eddolls', 'Full rebuild — Astra-format TIP with verified server inventory, corrected site mapping (Nashville→Atlanta, Lenexa→Boston), dev consolidation architecture, network/IP scheme, per-phase risk and validation sections', 'April 20, 2026'],
        ],
        col_widths=[0.4, 1.5, 3.8, 0.9]
    )

    # -----------------------------------------------------------------------
    # Implementation Summary
    # -----------------------------------------------------------------------
    add_section_heading(doc, 'Implementation Summary', 1)
    add_body(doc,
        'DialConnection operates a SIP/hosted-PBX platform that is customer-facing and revenue-critical. '
        'The SIP Gateways and PBX servers in Chicago, Dallas, Nashville, and Lenexa cannot tolerate '
        'extended unplanned downtime. Every phase in this plan is designed to keep replication running '
        'continuously until cutover, minimize cutover windows, and validate application-layer '
        'functionality before each site is declared live.'
    )
    add_body(doc,
        'The recommended cutover sequence is Atlanta and Boston first (smaller, lower-risk), then '
        'Chicago and Dallas. This allows the team to validate the Carbonite cutover process on '
        'smaller footprints before executing on the primary production sites. Active Directory '
        'domain controllers must be migrated before domain-joined servers at each site. SQL Availability '
        'Group configuration must be confirmed before the CST SQL cluster cutover in Chicago.'
    )

    # -----------------------------------------------------------------------
    # Timeline of Phases
    # -----------------------------------------------------------------------
    add_section_heading(doc, 'Timeline of Phases', 1)
    add_body(doc,
        'The following is the standard migration phase sequence. Dates are relative to project kickoff. '
        'Actual duration depends on DialConnection information-gathering completeness, '
        'maintenance window availability, and replication sync performance.'
    )
    add_table(doc,
        ['Phase', 'Name', 'Timing', 'Key Dependencies', 'Notes'],
        [
            ['Phase 1', 'Discovery and Environment Build',
             'Weeks 1–3', 'Kickoff call, DialConnection access, VM spec approval',
             'VDCs provisioned and all VMs built. No data migration yet.'],
            ['Phase 2', 'Networking, Security, and Replication Setup',
             'Weeks 3–5', 'Phase 1 VDCs and VMs complete',
             'VPNs, firewalls, MDR, EPP, patching agents, Carbonite replication initiated.'],
            ['Phase 3', 'Replication Monitoring and Pre-Cutover Validation',
             'Weeks 5–8', 'Initial sync complete, delta sync stable',
             'Backup and DRaaS configured and validated. Client validates apps in Thrive environment.'],
            ['Phase 4', 'Cutover',
             'Weeks 8–10', 'Phase 3 exit criteria met, windows approved',
             'Recommended order: Atlanta/Boston first, then Chicago/Dallas.'],
            ['Phase 5', 'Post-Cutover Validation and Handoff',
             'Weeks 10–12', 'All sites cut over and Availability Notifications issued',
             'Billing activation, $207K credit applied, DR test scheduled, steady-state handoff.'],
        ],
        col_widths=[0.65, 1.8, 1.2, 1.6, 1.35]
    )

    # -----------------------------------------------------------------------
    # 1. Project Overview
    # -----------------------------------------------------------------------
    add_section_heading(doc, '1. Project Overview', 1)
    add_body(doc,
        'This Technical Implementation Plan defines the approach, architecture, and phased execution '
        'strategy for migrating DialConnection LLC from its current TierPoint-hosted infrastructure to '
        'the Thrive ThriveCloud platform. DialConnection is a hosted PBX and VoIP provider whose '
        'SIP gateways, PBX platforms, and application servers are customer-facing and require careful '
        'scheduling of all cutover activities to avoid service impact.'
    )
    add_body(doc,
        'The migration is governed by Service Order CON-094865 (Q-69941), executed March 27, 2026. '
        'The service order covers 4 source datacenter locations (TierPoint Chicago, Dallas, Nashville, '
        'and Lenexa) migrating to 4 Thrive ThriveCloud Virtual Datacenters (Chicago, Dallas, Atlanta, '
        'and Boston). Note: Nashville maps to Thrive Atlanta and Lenexa maps to Thrive Boston — '
        'this corrects a labeling error in the original TIP documentation.'
    )
    add_bullet(doc, 'Client: ', 'DialConnection LLC')
    add_bullet(doc, 'Primary Contact: ', 'Keith Larson — President and COO')
    add_bullet(doc, 'Client Address: ', '309 Fellowship Road, Mount Laurel Township, NJ 08054')
    add_bullet(doc, 'Service Order: ', 'CON-094865 / Q-69941 — March 27, 2026')
    add_bullet(doc, 'Master Services Agreement: ', 'March 27, 2026')
    add_bullet(doc, 'Contract Term: ', '63 months (3 free months applied at billing activation)')
    add_bullet(doc, 'Thrive Project Manager: ', 'To be assigned at project kickoff')
    add_bullet(doc, 'Thrive Technical Lead: ', 'To be assigned at project kickoff')
    add_bullet(doc, 'Migration Source: ', 'TierPoint — Chicago, Dallas, Nashville, Lenexa')
    add_bullet(doc, 'Migration Target: ', 'ThriveCloud — Chicago, Dallas, Atlanta, Boston')
    add_bullet(doc, 'Migration Tool: ', 'Carbonite Migrate Standard (Windows and Linux) — 90 licenses contracted')
    add_bullet(doc, 'Active Directory Domain: ', 'dccontactcenter.com')

    # -----------------------------------------------------------------------
    # 2. Site Mapping
    # -----------------------------------------------------------------------
    add_section_heading(doc, '2. Site Mapping', 1)
    add_body(doc,
        'Four TierPoint source datacenter locations map to four Thrive ThriveCloud target Virtual '
        'Datacenters. The Nashville-to-Atlanta and Lenexa-to-Boston mappings were corrected from '
        'the original documentation and are authoritative as of April 14, 2026.'
    )
    add_table(doc,
        ['TierPoint Source', 'Thrive Target', 'Type', 'DR Pair', 'Total VMs', 'Windows', 'Linux', 'vCPU', 'RAM (GB)', 'Disk (GB)'],
        [
            ['Chicago', 'Chicago (ThriveCloud)', 'Dedicated Host', '↔ Dallas', '37', '29', '8', '247', '884', '21,173'],
            ['Dallas', 'Dallas (ThriveCloud)', 'Dedicated Host', '↔ Chicago', '51', '37', '14', '237', '854', '21,382'],
            ['Nashville', 'Atlanta (ThriveCloud)', 'Managed VDC', '↔ Boston', '4', '1', '3', '22', '26', '686'],
            ['Lenexa', 'Boston (ThriveCloud)', 'Managed VDC', '↔ Atlanta', '3', '1', '2', '15', '18', '788'],
            ['TOTAL', '', '', '', '95', '68', '27', '521', '1,782', '44,029'],
        ],
        col_widths=[0.9, 1.2, 1.0, 0.75, 0.65, 0.6, 0.5, 0.5, 0.75, 0.65],
        font_size=9
    )
    add_sub_note(doc,
        'VM counts reflect the April 13, 2026 inventory. The Service Order was contracted at 90 Carbonite '
        'licenses (90 VM migrations). Platform appliances (Alert Logic, FortiGate physical) and templates '
        'are not migrated via Carbonite. Dev server totals include 21 servers from Dallas and 1 from Chicago '
        'consolidating to Thrive Atlanta — these are counted within the Dallas and Chicago site totals above.'
    )

    # -----------------------------------------------------------------------
    # 3. Project Goals and Success Criteria
    # -----------------------------------------------------------------------
    add_section_heading(doc, '3. Project Goals and Success Criteria', 1)

    add_section_heading(doc, '3.1 Project Goals', 2)
    add_table(doc,
        ['#', 'Goal', 'Description'],
        [
            ['1', 'Deploy 4 Virtual Datacenters',
             'Provision and configure ThriveCloud VDCs at Atlanta, Boston, Chicago, and Dallas with all contracted compute, storage, networking, and firewall resources.'],
            ['2', 'Establish Server Backups',
             'Configure Thrive Managed Backup and Restore for all in-scope servers. Local backup plus offsite immutable backup at a secondary Thrive location for each site.'],
            ['3', 'Establish DRaaS Replication',
             'Configure DRaaS replication from each production site to its designated DR partner: Chicago↔Dallas, Atlanta↔Boston.'],
            ['4', 'Configure Security',
             'Deploy MDR Essentials SIEM collectors, configure FortiGate firewall security services, and install SentinelOne EPP agents on all in-scope servers.'],
            ['5', 'Replicate from TierPoint to Thrive',
             'Execute block-level server replication from TierPoint source environments to ThriveCloud via Carbonite Migrate for all in-scope servers.'],
            ['6', 'Cutover',
             'Execute per-site cutovers in approved maintenance windows. Validate all services before each Availability Notification is issued.'],
        ],
        col_widths=[0.3, 1.7, 4.6]
    )

    add_section_heading(doc, '3.2 Billing Activation Criteria', 2)
    add_body(doc,
        'Per the Service Order, Thrive will issue an Availability Notification when services at each '
        'location are fully provisioned and available. Within five (5) business days, DialConnection '
        'will review, test, and confirm each location is functioning as expected or provide a written '
        'summary of identified issues. Billing activates upon resolution of reported issues or '
        'thirty (30) days after the Availability Notification — whichever is earlier.'
    )
    add_sub_note(doc,
        'Two of five Dedicated Hosts at both Dallas and Chicago are placed in maintenance mode and reserved '
        'for DR failover. These hosts are not available for production workloads. The one-time credit of '
        '$207,165.60 (approximately 3 months MRC) is applied at billing activation and is non-transferable.'
    )

    # -----------------------------------------------------------------------
    # 4. Infrastructure Architecture
    # -----------------------------------------------------------------------
    add_section_heading(doc, '4. Infrastructure Architecture', 1)

    add_section_heading(doc, '4.1 Chicago Virtual Datacenter (Primary)', 2)
    add_body(doc,
        'Chicago is one of two primary production environments running on five ThriveCloud Dedicated '
        'Hosts (1.5TB RAM each, 144 licensed VMware cores). Two of the five hosts are reserved in '
        'maintenance mode for DR failover capacity. Thrive manages perimeter security via two '
        'FortiGate 200F appliances in HA configuration, replacing the current TierPoint FortiGate '
        '100F shared pair. Chicago replicates to Dallas for disaster recovery.'
    )
    add_table(doc,
        ['Resource', 'Contracted (Service Order)', 'Notes'],
        [
            ['Dedicated Hosts', '5 x ThriveCloud Dedicated Host (1.5TB RAM)', '2 of 5 in maintenance mode for DR.'],
            ['VMware Licensing', '144-core Dedicated Host VMware licenses', 'Per-core licensing for dedicated environment.'],
            ['Managed Windows VMs', '29 servers', 'Includes SQL licensing, Windows DC licensing, 24x7 monitoring.'],
            ['Managed Linux VMs', '6 servers', '24x7 monitoring and config management.'],
            ['Cloud Storage', '23 TB', 'Provisioned across virtual server volumes.'],
            ['SQL Licensing', '3 x 4-core SQL Std + 16 x 2-core SQL Std', 'For CSQLCST1/2/3 cluster.'],
            ['Windows DC Licensing', '72 x 2-core Windows Server DC', 'Datacenter edition licensing.'],
            ['Firewall', '2 x FortiGate 200F HA (24x7 managed)', 'IPS, AV/AS, HA service add-ons included.'],
            ['Public IPs', '256 Public IP Addresses', '1 base + 255 additional.'],
            ['Bandwidth', '200Mbps burstable to 1Gbps', ''],
            ['IPSEC VPN Tunnels', '16 Managed Site-to-Site VPN Tunnels', ''],
            ['Remote Access VPN', '100 users (FortiClient)', ''],
            ['Backup', '35 servers — local (46TB) + offsite immutable at Dallas (46TB)', ''],
            ['DRaaS', '28 servers replicated to Dallas (23TB at Dallas)', 'Includes DR VDC + annual DR test.'],
        ],
        col_widths=[2.0, 2.4, 2.2]
    )

    add_section_heading(doc, '4.2 Dallas Virtual Datacenter (Primary)', 2)
    add_body(doc,
        'Dallas is the second primary production environment, matching Chicago in dedicated host '
        'architecture. Dallas serves as the DR target for Chicago and hosts offsite immutable backup '
        'storage for Atlanta and Boston. Two FortiGate 200F appliances in HA manage perimeter '
        'security. Dallas replicates to Chicago for disaster recovery.'
    )
    add_table(doc,
        ['Resource', 'Contracted (Service Order)', 'Notes'],
        [
            ['Dedicated Hosts', '5 x ThriveCloud Dedicated Host (1.5TB RAM)', '2 of 5 in maintenance mode for DR.'],
            ['VMware Licensing', '144-core Dedicated Host VMware licenses', ''],
            ['Managed Windows VMs', '22 servers (SO contracted)', 'Note: inventory confirms 37W/14L total incl. dev servers — Change Log reconciliation pending.'],
            ['Managed Linux VMs', '26 servers (SO contracted)', 'See note above.'],
            ['Cloud Storage', '23 TB', ''],
            ['SQL Licensing', '2 x 4-core SQL Std + 12 x 2-core SQL Std', 'For EST (DSQLEST1/2) and PST (DSQLPST1).'],
            ['Windows DC Licensing', '72 x 2-core Windows Server DC', ''],
            ['Firewall', '2 x FortiGate 200F HA (24x7 managed)', 'IPS, AV/AS, HA service add-ons included.'],
            ['Public IPs', '256 Public IP Addresses', ''],
            ['Bandwidth', '200Mbps burstable to 1Gbps', ''],
            ['IPSEC VPN Tunnels', '18 Managed Site-to-Site VPN Tunnels', ''],
            ['Remote Access VPN', '100 users (FortiClient)', ''],
            ['Backup', '48 servers — local (46TB) + offsite immutable at Boston (46TB)', 'Hosts offsite backup for Atlanta and Boston.'],
            ['DRaaS', '28 servers replicated to Chicago (23TB at Chicago)', 'Includes DR VDC at Chicago.'],
        ],
        col_widths=[2.0, 2.4, 2.2]
    )

    add_section_heading(doc, '4.3 Atlanta Virtual Datacenter (Managed VDC)', 2)
    add_body(doc,
        'Atlanta is a Thrive-managed shared VDC environment hosting the Nashville production servers '
        '(Shared PBX for CST clients) plus the consolidated dev/QA environment from Dallas and Chicago. '
        'It uses an HA virtual firewall with IPS and DDoS protection built into the VDC base service '
        '(no physical FortiGate). Atlanta replicates to Boston for disaster recovery. '
        'Offsite immutable backup is stored at Dallas.'
    )
    add_table(doc,
        ['Resource', 'Contracted (Service Order)', 'Notes'],
        [
            ['Virtual Datacenter', 'ThriveCloud VDC — Atlanta', 'HA virtual firewall, IPS, DDoS protection, 1 public IP base.'],
            ['Managed Windows VMs', '1 server', 'DCloud3NDRSQL1 (DR SQL).'],
            ['Managed Linux VMs', '2 servers', 'DCloud3NPBX01 (PBX) + NSH-DRGateway-PBX.'],
            ['Cloud Storage', '2 TB', ''],
            ['Compute', '9 x 2GB vRAM blocks (1 vCPU each) + 6 additional vCPUs', ''],
            ['Public IPs', '5 additional (6 total including base)', ''],
            ['Bandwidth', '50Mbps burstable to 150Mbps', ''],
            ['Remote Access VPN', '100 users', ''],
            ['Backup', '3 servers — local (4TB) + offsite immutable at Dallas (4TB)', ''],
            ['DRaaS', '3 servers replicated to Boston (4TB at Boston)', 'Includes DR VDC at Boston + annual DR test.'],
        ],
        col_widths=[2.0, 2.4, 2.2]
    )
    add_sub_note(doc,
        'Dev/QA Consolidation: 22 dev servers (21 from Dallas + DCloud3CSQLDev1 from Chicago) are being '
        'consolidated to Thrive Atlanta as a deliberate architectural decision. These servers are covered '
        'under the Dallas and Chicago managed server licensing and are not separate Atlanta VDC managed '
        'VMs. The Atlanta VDC compute resources are sized to accommodate the production PBX workload. '
        'Dev server sizing will be confirmed during VM specification review in Phase 1.'
    )

    add_section_heading(doc, '4.4 Boston Virtual Datacenter (Managed VDC)', 2)
    add_body(doc,
        'Boston is a Thrive-managed shared VDC environment hosting the Lenexa production servers '
        '(Shared PBX for EST and PST clients). It mirrors the Atlanta VDC architecture with a managed '
        'virtual firewall. Boston replicates to Atlanta for disaster recovery. Offsite immutable '
        'backup is stored at Dallas.'
    )
    add_table(doc,
        ['Resource', 'Contracted (Service Order)', 'Notes'],
        [
            ['Virtual Datacenter', 'ThriveCloud VDC — Boston', 'HA virtual firewall, IPS, DDoS protection, 1 public IP base.'],
            ['Managed Windows VMs', '1 server', 'DCloud3LDRSQL1 (DR SQL).'],
            ['Managed Linux VMs', '3 servers', 'DCloud3KPBX01 (PBX) + LEN-DRGateway-PBX. Note: SO contracts 3 Linux — confirm third server.'],
            ['Cloud Storage', '2 TB', ''],
            ['Compute', '13 x 2GB vRAM blocks (1 vCPU each) + 9 additional vCPUs', ''],
            ['Public IPs', '5 additional (6 total including base)', ''],
            ['Bandwidth', '50Mbps burstable to 150Mbps', ''],
            ['Remote Access VPN', '100 users', ''],
            ['Backup', '3 servers — local (4TB) + offsite immutable at Dallas (4TB)', ''],
            ['DRaaS', '3 servers replicated to Atlanta (4TB at Atlanta)', 'Includes DR VDC at Atlanta.'],
        ],
        col_widths=[2.0, 2.4, 2.2]
    )

    add_section_heading(doc, '4.5 Cross-Site DR and Replication Architecture', 2)
    add_table(doc,
        ['Production Site', 'DR Site', 'Servers in Replication', 'DR Storage', 'SLA', 'DR VDC'],
        [
            ['Chicago', 'Dallas', '28 servers', '23 TB at Dallas', '15-min failover init', 'Thrive Cloud DR VDC — Dallas'],
            ['Dallas', 'Chicago', '28 servers', '23 TB at Chicago', '15-min failover init', 'Thrive Cloud DR VDC — Chicago'],
            ['Atlanta', 'Boston', '3 servers', '4 TB at Boston', '15-min failover init', 'Thrive Cloud DR VDC — Boston'],
            ['Boston', 'Atlanta', '3 servers', '4 TB at Atlanta', '15-min failover init', 'Thrive Cloud DR VDC — Atlanta'],
        ],
        col_widths=[0.9, 0.9, 1.35, 1.0, 1.2, 1.25]
    )
    add_sub_note(doc,
        'DRaaS SLA: Thrive initiates failover within 15 minutes of a confirmed disaster declaration. '
        'Recovery target: up to 40 virtual servers within 1 hour, then 10 additional every 15 minutes. '
        'RPO < 15 minutes for all protected servers. Annual user-participated DR test included. '
        'Client must run simulated recovery every 12 months for SLA to remain in effect.'
    )

    # -----------------------------------------------------------------------
    # 5. Server Inventory
    # -----------------------------------------------------------------------
    add_section_heading(doc, '5. Server Inventory', 1)
    add_body(doc,
        'The following tables list all virtual servers and platform infrastructure across all four '
        'TierPoint source locations as confirmed in the April 2026 inventory. Yellow fields in the '
        'Information Gathering Workbook require DialConnection input before Phase 1 can close. '
        '"Migrate" column: Y = migrate via Carbonite, N = decommission/replace at Thrive.'
    )

    add_section_heading(doc, '5.1 Chicago (→ Thrive Chicago)', 2)
    add_sub_note(doc, 'Total: 37 VMs (29 Windows, 8 Linux) + 2 physical FortiGates. DCloud3CSQLDev1 routes to Thrive Atlanta (dev consolidation). Alert Logic appliances and physical FortiGates are not migrated.')
    add_table(doc, SERVER_HEADERS, CHICAGO_SERVERS, SERVER_WIDTHS, font_size=9)

    add_section_heading(doc, '5.2 Dallas (→ Thrive Dallas) — Production Servers', 2)
    add_sub_note(doc, 'Total Dallas: 51 VMs (37 Windows, 14 Linux) + 2 physical FortiGates. Production servers below. Dev/QA servers consolidating to Atlanta are in Section 5.3.')
    add_table(doc, SERVER_HEADERS, DALLAS_SERVERS, SERVER_WIDTHS, font_size=9)

    add_section_heading(doc, '5.3 Dallas Dev/QA Environment (→ Thrive Atlanta)', 2)
    add_sub_note(doc, 'These 21 servers are physically located at TierPoint Dallas but are migrating to Thrive Atlanta as part of the dev environment consolidation. DC3VBVM03/04 require existence confirmation from DialConnection.')
    add_table(doc, SERVER_HEADERS, DALLAS_DEV_TO_ATL, SERVER_WIDTHS, font_size=9)

    add_section_heading(doc, '5.4 Nashville (→ Thrive Atlanta)', 2)
    add_sub_note(doc, 'Total: 4 VMs (1 Windows, 3 Linux) + 1 physical FortiGate. CORRECTED: Nashville maps to Thrive Atlanta (not Boston).')
    add_table(doc, SERVER_HEADERS, NASHVILLE_ATL_SERVERS, SERVER_WIDTHS, font_size=9)

    add_section_heading(doc, '5.5 Lenexa (→ Thrive Boston)', 2)
    add_sub_note(doc, 'Total: 3 VMs (1 Windows, 2 Linux) + 1 physical FortiGate. CORRECTED: Lenexa maps to Thrive Boston (not Atlanta). Service Order contracts 3 Linux VMs — third server requires confirmation from DialConnection.')
    add_table(doc, SERVER_HEADERS, LENEXA_BOS_SERVERS, SERVER_WIDTHS, font_size=9)

    # -----------------------------------------------------------------------
    # 6. Network and IP Architecture
    # -----------------------------------------------------------------------
    add_section_heading(doc, '6. Network and IP Architecture', 1)

    add_section_heading(doc, '6.1 Site IP Summary', 2)
    add_table(doc,
        ['Site', 'TierPoint Private', 'TierPoint Public', 'TierPoint VPN', 'Thrive Allocated', 'Gateway', 'Primary DNS', 'Secondary DNS'],
        [
            ['Chicago', '10.5.1.0/24', '66.36.23.0/25', '66.36.22.128/25', '216.38.67.0/24', '10.5.1.1', '10.5.1.10 (CPDC)', '10.5.1.11 (CPDC1)'],
            ['Dallas', '10.7.1.0/24', '40.142.107.0/25', '40.142.106.128/25', '74.220.95.0/24', '10.7.1.1', '10.7.1.10 (DPDC)', '10.7.1.11 (DSDC1)'],
            ['Nashville→Atlanta', '10.6.1.0/24', '40.143.148.48/29', 'N/A', 'N/A', '10.6.1.1', '10.6.1.10', '10.6.1.11'],
            ['Lenexa→Boston', '10.8.1.0/24', '40.142.82.80/29', 'N/A', 'N/A', '10.8.1.1', '10.8.1.10', '(single DNS)'],
        ],
        col_widths=[1.0, 0.95, 0.95, 0.95, 0.95, 0.8, 0.75, 0.95],
        font_size=9
    )
    add_body(doc,
        'DNS upstream for all sites (FortiGate): 96.45.45.45 / 96.45.46.46 (DNS over TLS). '
        'Active Directory domain: dccontactcenter.com across all sites.'
    )

    add_section_heading(doc, '6.2 Key Network Segments by Site', 2)
    add_table(doc,
        ['Site', 'VLAN / Name', 'Subnet', 'Purpose', 'Notes'],
        [
            ['Chicago', 'VLAN 3044 / PROD-Private-CHI', '10.5.1.0/24', 'Production LAN', 'PDC 10.5.1.10/.11'],
            ['Chicago', 'VLAN 3046 / CHI-VPN', '66.36.22.128/25', 'IPSEC VPN Public Block', 'Endpoint for all S2S tunnels'],
            ['Chicago', 'VLAN 3050 / CHI-Public', '40.142.29.64/27', 'Public Services (PBX/RDS)', 'Inbound SIP, RDS, app servers'],
            ['Chicago', 'VLAN 3045 / DR-Private-DA2', '10.7.1.0/24', 'DR Network (Dallas VMs)', 'Mirrors Dallas prod LAN'],
            ['Dallas', 'VLAN 3050 / PRD-Private-DA2', '10.7.1.0/24', 'Production LAN', 'PDC 10.7.1.10/.11'],
            ['Dallas', 'VLAN 3056 / DA2-VPN', '40.142.106.128/25', 'IPSEC VPN Public Block', 'Endpoint for all S2S tunnels'],
            ['Dallas', 'VLAN 3057 / DA2-Public', '40.142.109.192/27', 'Public Services (PBX/RDS)', 'Inbound SIP, RDS, app servers'],
            ['Dallas', 'VLAN 3123 / DEV-Private-DA2', '10.6.1.0/24', 'DEV Environment LAN', 'All dev servers → Atlanta VDC'],
            ['Nashville→ATL', 'VLAN 3022 / PRD-Private', '10.6.1.0/24', 'Production LAN (PBX)', 'PBX contact center'],
            ['Nashville→ATL', 'VLAN 3023 / PBX-Internal', '10.6.2.0/24', 'PBX Internal VLAN', 'UNDOCUMENTED — plan carefully'],
            ['Nashville→ATL', 'VLAN 3128 / PROD-Public_2', '40.143.108.16/29', 'Secondary Public Block', 'UNDOCUMENTED — confirm w/ client'],
            ['Lenexa→BOS', 'VLAN 3034 / PRD-Private-LEN', '10.8.1.0/24', 'Production LAN (PBX)', 'PBX contact center'],
            ['Lenexa→BOS', 'VLAN 3035 / Direct IP', '40.142.82.80/29', 'Public Services (PBX/SIP)', 'Inbound PBX, SIP, TCP/34225'],
        ],
        col_widths=[1.0, 1.7, 1.0, 1.35, 1.55],
        font_size=9
    )
    add_sub_note(doc,
        'Open items: Two undocumented VLANs at Nashville (VLAN 3023 / PBX-Internal and VLAN 3128 / PROD-Public_2) '
        'must be documented by DialConnection before Nashville/Atlanta cutover planning can be finalized. '
        'Full public IP usage mapping (what each public IP is assigned to) is required from DialConnection '
        'for both Chicago and Dallas before cutover.'
    )

    # -----------------------------------------------------------------------
    # 7. Backup Architecture
    # -----------------------------------------------------------------------
    add_section_heading(doc, '7. Backup Architecture', 1)
    add_body(doc,
        'Thrive Managed Backup and Restore Cloud is configured for all in-scope servers across all '
        'four datacenters. Each site maintains a local backup repository for fast recovery, plus a '
        'daily offsite immutable copy to a geographically separate Thrive location to protect against '
        'site-level failures and ransomware. 30-day data retention is enforced. All backup restore '
        'requests are handled via the Thrive Client Portal.'
    )
    add_table(doc,
        ['Site', 'Servers Backed Up', 'Local Backup Storage', 'Offsite Immutable Location', 'Offsite Storage'],
        [
            ['Chicago', '35 (contracted)', '46 TB at Chicago', 'Dallas', '46 TB'],
            ['Dallas', '48 (contracted)', '46 TB at Dallas', 'Boston', '46 TB'],
            ['Atlanta', '3 (contracted)', '4 TB at Atlanta', 'Dallas', '4 TB'],
            ['Boston', '3 (contracted)', '4 TB at Boston', 'Dallas', '4 TB'],
        ],
        col_widths=[0.9, 1.3, 1.6, 1.6, 1.2]
    )
    add_sub_note(doc,
        'Backup counts are per the Service Order contracted quantities. Updated inventory (95 VMs) '
        'may require a Change Log adjustment. Application-level recovery from a backup restore is '
        'DialConnection\'s responsibility. DialConnection must notify Thrive of any server additions '
        'or removals from backup scope.'
    )

    # -----------------------------------------------------------------------
    # 8. Security Architecture
    # -----------------------------------------------------------------------
    add_section_heading(doc, '8. Security Architecture', 1)

    add_section_heading(doc, '8.1 Managed Detection and Response (MDR)', 2)
    add_body(doc,
        'Thrive MDR Essentials provides 24x7x365 security monitoring through a dedicated SIEM tenant '
        'with two virtual collector VMs deployed at Dallas and Chicago. All security events are '
        'triaged by the Thrive Security Operations Center. This service covers 98 infrastructure '
        'and application devices plus 15 Microsoft 365 mailboxes.'
    )
    add_table(doc,
        ['MDR Component', 'Detail', 'Notes'],
        [
            ['MDR Base Service', 'Thrive MDR Essentials — dedicated tenant', '24x7x365 SOC monitoring, alerting, and incident response.'],
            ['Cloud Collectors', '2 x MDR Collector VMs — Dallas + Chicago', '8 vCPU, 8GB vRAM, 125GB storage each. VPN connection to remote sites.'],
            ['Device Coverage', '98 infrastructure / application devices', '10 EPS allowance per device.'],
            ['M365 Monitoring', '15 mailboxes', 'Office 365 log collection and monitoring.'],
            ['Log Retention', '1-year extended log retention (98 devices)', ''],
            ['High Severity Response', '15-minute initial response', 'Phone + incident ticket. Criteria: likely breach in progress.'],
            ['Medium Severity Response', '4-hour initial response', 'Incident ticket. Criteria: blocked attempts or abnormal behavior.'],
            ['Low Severity Response', '24-hour initial response', 'Incident ticket. Criteria: single non-impacting event.'],
        ],
        col_widths=[1.9, 2.3, 2.4]
    )

    add_section_heading(doc, '8.2 Managed Enterprise Endpoint Protection (SentinelOne)', 2)
    add_body(doc,
        'Thrive Managed Enterprise Endpoint Security and Response (SentinelOne Complete) is deployed '
        'to all 90 contracted servers. The service provides behavioral AI-based threat detection, '
        'automated remediation, device isolation, and 24x7 SOC alerting.'
    )
    add_bullet(doc, 'Coverage: ', '90 servers — Complete tier (per server)')
    add_bullet(doc, 'Detection: ', 'Behavioral AI, fileless attack detection, MITRE ATT&CK threat hunting')
    add_bullet(doc, 'Response: ', 'Automated remediation, device isolation on infection or attack detection')
    add_bullet(doc, 'Validation: ', 'Quarterly Security Control Validation exercise on a designated test server')

    add_section_heading(doc, '8.3 Perimeter Firewall Services', 2)
    add_table(doc,
        ['Service Component', 'Chicago', 'Dallas', 'Atlanta / Boston'],
        [
            ['Firewall Platform', '2 x FortiGate 200F HA', '2 x FortiGate 200F HA', 'Managed Virtual Firewall (VDC)'],
            ['IPS', 'Managed Firewall IPS', 'Managed Firewall IPS', 'Included in VDC base'],
            ['Antivirus / Anti-Spyware', 'Managed AV / Anti-Spyware', 'Managed AV / Anti-Spyware', 'Included in VDC base'],
            ['High Availability', 'Managed HA Service', 'Managed HA Service', 'Included in VDC base'],
            ['IPSEC Site-to-Site VPN', '16 managed tunnels', '18 managed tunnels', 'N/A'],
            ['Remote Access VPN', '100 users (FortiClient)', '100 users (FortiClient)', '100 users each'],
            ['DDoS Protection', 'Included in VDC baseline', 'Included in VDC baseline', 'Included in VDC base'],
        ],
        col_widths=[1.8, 1.6, 1.6, 1.6]
    )

    # -----------------------------------------------------------------------
    # 9. Managed Server Patching
    # -----------------------------------------------------------------------
    add_section_heading(doc, '9. Managed Server Patching', 1)
    add_body(doc,
        'All 90 contracted servers are enrolled in Thrive Managed Server Patching Enterprise Plus '
        'with Third-Party Advanced Patching. Thrive deploys a Kaseya RMM agent on each server to '
        'automate patch scanning, deployment, and remediation. Patching schedules are configured '
        'collaboratively with DialConnection.'
    )
    add_table(doc,
        ['Feature', 'Detail'],
        [
            ['Server Coverage', '90 servers — Enterprise Plus + Third-Party Advanced Patching'],
            ['OS Patching SLA', 'Critical and security patches deployed within 30 days of vendor release'],
            ['Third-Party Patching', 'Approved third-party application patching via Advanced Patching service'],
            ['Patch Scanning', 'Lightweight patch scan every 3 days to identify missing patches'],
            ['Remediation', 'Automated re-deployment up to 3 attempts; manual remediation if re-deployments fail'],
            ['Patching Exclusions', 'OS version upgrades, BIOS/driver updates, and Exchange/SQL updates from non-Windows Update sources'],
        ],
        col_widths=[2.2, 4.4]
    )

    # -----------------------------------------------------------------------
    # 10. RDS Licensing
    # -----------------------------------------------------------------------
    add_section_heading(doc, '10. Remote Desktop Services Licensing', 1)
    add_body(doc,
        'Thrive provides 50 Windows RDS Access Licenses (per user) under the contracted services. '
        'These licenses support remote desktop access to Windows Server sessions. DialConnection '
        'may reduce or cancel this line item with 30 days\' prior written notice to Thrive, '
        'effective on the last day of the following month.'
    )
    add_bullet(doc, 'Quantity: ', '50 user RDS Access Licenses (HSSPLA003 — per user)')
    add_bullet(doc, 'Sites: ', 'RDS farms at Chicago and Dallas — Public farm (whitelist access) + standalone internal RDS at both sites')
    add_bullet(doc, 'Cancellation: ', '30 days written notice required')

    # -----------------------------------------------------------------------
    # 11. Migration Strategy
    # -----------------------------------------------------------------------
    add_section_heading(doc, '11. Migration Strategy', 1)

    add_section_heading(doc, '11.1 Migration Approach', 2)
    add_body(doc,
        'Server migration from TierPoint to ThriveCloud is executed via block-level replication using '
        'Carbonite Migrate Standard for Windows and Linux. This allows continuous synchronization '
        'of server data prior to cutover, minimizing downtime to a short final-sync window. '
        '90 licenses are contracted; Change Log reconciliation against the 95-VM inventory is in progress.'
    )
    add_table(doc,
        ['Step', 'Activity', 'Notes'],
        [
            ['1 — Pre-Migration', 'Thrive provisions all 4 ThriveCloud VDCs, configures networking and firewalls', 'VDCs must be fully operational before replication begins.'],
            ['2 — Agent Install', 'Carbonite Migrate agents installed on all in-scope source servers at TierPoint', 'Full admin access to source servers and hypervisors required.'],
            ['3 — Initial Sync', 'Block-level replication of full server disk images from TierPoint to ThriveCloud', 'Initial sync duration varies; large SQL servers (CSQLCST1 ~6.9TB) will take several days.'],
            ['4 — Delta Sync', 'Continuous incremental replication of changed blocks until cutover', 'Delta sync keeps target servers within minutes of source.'],
            ['5 — Pre-Cutover Test', 'DialConnection validates application and service functionality in ThriveCloud environment', 'Must be complete before cutover window is scheduled.'],
            ['6 — Cutover', 'Final sync, server re-IP, DNS/routing changes in approved maintenance window', 'Thrive re-IPs servers post-migration where applicable.'],
            ['7 — Post-Cutover', 'Backup and DRaaS validated; TierPoint decommission initiated by DialConnection', 'Thrive is not responsible for TierPoint decommission.'],
        ],
        col_widths=[1.2, 2.4, 3.0]
    )

    add_section_heading(doc, '11.2 Migration Scope by Site', 2)
    add_table(doc,
        ['Site', 'Migrate To', 'Production VMs', 'Dev VMs', 'Not Migrating', 'Cutover Notes'],
        [
            ['Chicago', 'Thrive Chicago', '35 (29W + 6L gateways)', '1 (CSQLDev1 → Atlanta)', 'AL appliances (2), FortiGates (2)', 'AD and SQL servers must migrate first. Gateways last within site.'],
            ['Dallas', 'Thrive Dallas', '24 (22W + 2L gateways, 4 DR gateways)', '21 (→ Atlanta VDC)', 'AL appliances (2), FortiGates (2)', 'AD and SQL servers first. 4 DR gateways flagged for possible removal — confirm.'],
            ['Nashville', 'Thrive Atlanta', '4 (1W + 3L)', 'None', 'FortiGate 80F HA (1)', 'Smaller site — recommend cutting over before Chicago/Dallas.'],
            ['Lenexa', 'Thrive Boston', '3 (1W + 2L)', 'None', 'FortiGate 80F HA (1)', 'Smaller site — recommend cutting over before Chicago/Dallas.'],
        ],
        col_widths=[0.75, 0.95, 1.35, 1.1, 1.35, 1.1]
    )

    add_section_heading(doc, '11.3 Dev Environment Consolidation to Atlanta', 2)
    add_body(doc,
        'As a deliberate architectural decision, all development and QA servers are being consolidated '
        'to Thrive Atlanta during this migration rather than maintaining separate dev environments at '
        'Chicago and Dallas. This simplifies dev operations, centralizes the QA footprint, and avoids '
        'maintaining dev infrastructure on the Dedicated Host environments.'
    )
    add_bullet(doc, 'From Dallas (21 servers): ', 'DC3DevPDC, DC3DevDB1, DC3DevGW001–004, DC3DevQA001–007, DC3DevWeb1/2, DC3DGPXStaging, DC3DAppStaging, DC3VBVM01/02 (+ verify DC3VBVM03/04)')
    add_bullet(doc, 'From Chicago (1 server): ', 'DCloud3CSQLDev1')
    add_bullet(doc, 'Important: ', 'Two Windows XP legacy VMs (DC3VBVM01/02) and two unconfirmed VMs (DC3VBVM03/04) require DialConnection decision: migrate vs. rebuild. Windows XP is not supported for Carbonite migration and is outside Thrive DRaaS scope.')
    add_sub_note(doc,
        'The Atlanta VDC is sized per the Service Order for the Nashville production workload. '
        'Atlanta VDC compute resources for the dev consolidation will be assessed during Phase 1 '
        'VM specification review. Additional compute may require a Change Log adjustment.'
    )

    # -----------------------------------------------------------------------
    # 12. Implementation Phases
    # -----------------------------------------------------------------------
    add_section_heading(doc, '12. Implementation Phases', 1)
    add_body(doc,
        'The migration is organized into five sequential phases. No phase begins until the prior '
        'phase\'s exit criteria are met and documented. Thrive will work with DialConnection to avoid '
        'impacting customer-facing services during normal operating hours of 8:00 AM to 9:00 PM EST '
        'Monday through Friday.'
    )

    # --- Phase 1 ---
    add_section_heading(doc, 'Phase 1 — Discovery and Environment Build', 2)
    add_body(doc,
        'Phase 1 establishes the foundational Thrive infrastructure and confirms all specifications '
        'required to proceed with migration. No data is moved during this phase — the objective is '
        'a fully built, empty ThriveCloud environment ready to accept Carbonite replication. '
        'DialConnection must complete the Information Gathering Workbook before Phase 1 can close.'
    )
    add_section_heading(doc, 'Prerequisites', 3)
    add_bullet(doc, '', 'Project kickoff call scheduled and completed')
    add_bullet(doc, '', 'Thrive PM and Technical Lead assigned')
    add_bullet(doc, '', 'DialConnection designates Project POC and Technical/IT Lead')
    add_bullet(doc, '', 'DialConnection provides full admin access to all source servers and hypervisors')
    add_section_heading(doc, 'Key Activities', 3)
    add_checklist_item(doc, 'Project kickoff — confirm contacts, roles, schedule, and recurring meeting cadence')
    add_checklist_item(doc, 'DialConnection completes Information Gathering Workbook (all yellow fields)')
    add_checklist_item(doc, 'Thrive reviews and finalizes VM Specification list for all 95 servers (hostname, OS, vCPU, vRAM, storage)')
    add_checklist_item(doc, 'DialConnection reviews and approves VM Specification list')
    add_checklist_item(doc, 'Thrive provisions all 4 ThriveCloud VDCs (Chicago, Dallas, Atlanta, Boston)')
    add_checklist_item(doc, 'Thrive provisions all virtual machines per approved VM Spec list')
    add_checklist_item(doc, 'Thrive configures Dedicated Host environments at Chicago and Dallas (VMware, compute, storage)')
    add_checklist_item(doc, 'Thrive deploys and configures FortiGate 200F HA pairs at Chicago and Dallas')
    add_checklist_item(doc, 'Thrive configures virtual networks, IP addressing, and VLAN segmentation per approved specifications')
    add_checklist_item(doc, 'Dev consolidation compute sizing confirmed for Atlanta VDC — Change Log submitted if adjustment needed')
    add_section_heading(doc, 'Risks and Contingencies', 3)
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['DialConnection delays providing workbook data or server access', 'Medium', 'High', 'Thrive PM escalates to Keith Larson. Phase 1 timeline adjusted accordingly — delay is DialConnection-owned.'],
            ['VM Specification requires significant changes after approval', 'Low', 'Medium', 'Change Log submitted. Additional compute billed on prorated basis.'],
            ['Dev server sizing exceeds Atlanta VDC contracted compute', 'Medium', 'Medium', 'Atlanta VDC compute Change Log submitted. Sizing confirmed during VM Spec review.'],
            ['DC3VBVM03/04 exist and cannot be Carbonite-migrated (XP)', 'Medium', 'Low', 'DialConnection decides: rebuild at Thrive Atlanta or decommission. Confirm in Phase 1.'],
        ],
        col_widths=[1.9, 0.85, 0.7, 3.15]
    )
    add_section_heading(doc, 'Phase 1 Exit Criteria', 3)
    add_bullet(doc, '', 'All 4 ThriveCloud VDCs provisioned and all VMs built per approved VM Specification list')
    add_bullet(doc, '', 'FortiGate 200F HA pairs deployed and basic network connectivity confirmed at Chicago and Dallas')
    add_bullet(doc, '', 'DialConnection VM Specification approval documented in writing')
    add_bullet(doc, '', 'Information Gathering Workbook complete — no open yellow fields blocking Phase 2')

    # --- Phase 2 ---
    add_section_heading(doc, 'Phase 2 — Networking, Security, and Replication Setup', 2)
    add_body(doc,
        'Phase 2 establishes all network connectivity, security tooling, and initiates Carbonite '
        'replication from TierPoint to ThriveCloud. This is the most activity-dense phase — VPN '
        'tunnels, firewall policies, MDR collectors, EPP agents, patching agents, and replication '
        'can all be configured in parallel once the VDCs from Phase 1 are ready.'
    )
    add_section_heading(doc, 'Prerequisites', 3)
    add_bullet(doc, '', 'Phase 1 exit criteria met and documented')
    add_bullet(doc, '', 'DialConnection provides VPN tunnel inventory (all IPSEC tunnels at Chicago and Dallas)')
    add_bullet(doc, '', 'DialConnection provides firewall rule export from existing TierPoint FortiGates')
    add_bullet(doc, '', 'DialConnection provides full public IP usage mapping for Chicago and Dallas')
    add_bullet(doc, '', 'DialConnection provides device list and network diagram for MDR log collection scope')
    add_section_heading(doc, 'Key Activities', 3)
    add_checklist_item(doc, 'Thrive configures all IPSEC site-to-site VPN tunnels (16 at Chicago, 18 at Dallas)')
    add_checklist_item(doc, 'Thrive configures remote access VPN (FortiClient) at all four sites — 100 users each')
    add_checklist_item(doc, 'Thrive configures FortiGate IPS, antivirus/anti-spyware, and HA policies at Chicago and Dallas')
    add_checklist_item(doc, 'Thrive configures virtual firewall policies at Atlanta and Boston VDCs')
    add_checklist_item(doc, 'Thrive deploys MDR collector VMs at Dallas and Chicago (8 vCPU, 8GB, 125GB storage)')
    add_checklist_item(doc, 'Thrive configures MDR SIEM tenant and begins log collection from all applicable sources')
    add_checklist_item(doc, 'Thrive installs Carbonite Migrate agents on all in-scope source servers at TierPoint')
    add_checklist_item(doc, 'Initial block-level replication initiated from TierPoint to ThriveCloud for all sites')
    add_checklist_item(doc, 'Thrive installs SentinelOne EPP agents on all 90 contracted servers')
    add_checklist_item(doc, 'Thrive deploys Kaseya RMM patching agents and configures patching schedules with DialConnection')
    add_section_heading(doc, 'Risks and Contingencies', 3)
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['Missing VPN tunnel inventory delays firewall config', 'Medium', 'High', 'Thrive reviews existing TierPoint FortiGate config to gather data. Client must authorize access. Adds 3–5 days.'],
            ['Initial Carbonite sync is slower than expected for large SQL servers', 'Medium', 'Medium', 'SQL servers (CSQLCST1 ~6.9TB, DSQLEST1 ~4.1TB, DSQLPST1 ~3.4TB) will take longest. Begin replication immediately after Phase 2 starts. Monitor sync rate daily.'],
            ['Undocumented Nashville VLANs (3023, 3128) cause connectivity issues at Atlanta', 'Medium', 'High', 'DialConnection must document all Nashville VLANs before Nashville/Atlanta Phase 4 cutover. Firewall policy built conservatively.'],
            ['EPP agent ports blocked by existing firewall rules', 'Low', 'Medium', 'Thrive manages Chicago and Dallas firewalls — ports opened as needed. DialConnection escalation contact required for any app-layer conflicts.'],
        ],
        col_widths=[1.9, 0.85, 0.7, 3.15]
    )
    add_section_heading(doc, 'Phase 2 Exit Criteria', 3)
    add_bullet(doc, '', 'All VPN tunnels configured and validated at Chicago and Dallas')
    add_bullet(doc, '', 'Carbonite replication actively running and syncing for all in-scope servers at all 4 sites')
    add_bullet(doc, '', 'MDR log collection active and SIEM tenant receiving events')
    add_bullet(doc, '', 'EPP agents installed and reporting to Thrive SOC on all 90 contracted servers')
    add_bullet(doc, '', 'Patching agents deployed and first patch scan completed')

    # --- Phase 3 ---
    add_section_heading(doc, 'Phase 3 — Replication Monitoring and Pre-Cutover Validation', 2)
    add_body(doc,
        'Phase 3 runs concurrently with ongoing Carbonite delta sync. The objective is to validate '
        'that all services and applications function correctly in the ThriveCloud environment before '
        'any cutover is scheduled. Backup and DRaaS are configured and validated. DialConnection '
        'identifies and resolves any application compatibility or performance issues during this phase.'
    )
    add_section_heading(doc, 'Prerequisites', 3)
    add_bullet(doc, '', 'Phase 2 exit criteria met — all replication running, all agents deployed')
    add_bullet(doc, '', 'DialConnection designates maintenance/cutover windows for each site')
    add_bullet(doc, '', 'DialConnection provides DR recovery group dependencies and boot order')
    add_section_heading(doc, 'Key Activities', 3)
    add_checklist_item(doc, 'Thrive monitors Carbonite replication health and delta sync performance across all sites daily')
    add_checklist_item(doc, 'DialConnection validates application and service functionality in ThriveCloud environment for each site')
    add_checklist_item(doc, 'DialConnection confirms SIP Gateway connectivity and PBX functionality from Thrive environment')
    add_checklist_item(doc, 'DialConnection confirms SQL Availability Group / replication configuration in ThriveCloud')
    add_checklist_item(doc, 'DialConnection identifies and communicates any performance or compatibility issues to Thrive')
    add_checklist_item(doc, 'Thrive configures Managed Backup and Restore for all servers at all four sites')
    add_checklist_item(doc, 'Thrive validates local backup completion and offsite immutable copy jobs for all sites')
    add_checklist_item(doc, 'Thrive configures DRaaS replication for all sites (Chicago↔Dallas, Atlanta↔Boston)')
    add_checklist_item(doc, 'Thrive validates DRaaS replication health and RPO compliance across all pairs')
    add_checklist_item(doc, 'Cutover runbook built and reviewed by both parties before Phase 4 is authorized')
    add_checklist_item(doc, 'DialConnection approves cutover schedule for each site in writing')
    add_section_heading(doc, 'Risks and Contingencies', 3)
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['Application issues discovered during client validation delay cutover', 'Medium', 'High', 'Application compatibility issues are DialConnection responsibility per Service Order. Thrive assists with infrastructure-layer troubleshooting. SQL/application issues not in Thrive scope.'],
            ['DRaaS RPO cannot meet <15 min SLA due to insufficient bandwidth at TierPoint', 'Low', 'High', 'Client must provision sufficient bandwidth at source per DR SLA requirements. Thrive documents bandwidth requirement. SLA contingent on client-side compliance.'],
            ['SIP gateway IP-dependent configurations require manual re-IP work beyond Carbonite scope', 'Medium', 'High', 'IP and hostname variables are client responsibility per Service Order. DialConnection must identify all variables requiring update in Thrive environment during Phase 3 testing.'],
            ['Backup jobs for large SQL servers fail due to storage volume sizing', 'Low', 'Medium', 'Storage volumes confirmed during VM Spec (Phase 1). Thrive monitors backup jobs and auto-retries. Change Log for additional storage if required.'],
        ],
        col_widths=[1.9, 0.85, 0.7, 3.15]
    )
    add_section_heading(doc, 'Phase 3 Exit Criteria', 3)
    add_bullet(doc, '', 'DialConnection has validated application and service functionality in ThriveCloud for each site')
    add_bullet(doc, '', 'Backup and DRaaS operational and validated for all sites')
    add_bullet(doc, '', 'Carbonite delta sync stable (lag < 15 minutes) for all in-scope servers')
    add_bullet(doc, '', 'Cutover runbook reviewed and approved by both parties')
    add_bullet(doc, '', 'Maintenance windows confirmed for each site')

    # --- Phase 4 ---
    add_section_heading(doc, 'Phase 4 — Cutover', 2)
    add_body(doc,
        'Phase 4 executes the production cutover for each site in the approved maintenance windows. '
        'The recommended order is Atlanta and Boston first (smaller footprint, lower risk), then '
        'Chicago and Dallas. Each site cutover is independent — a successful Atlanta cutover does '
        'not require Boston to cut over simultaneously.'
    )
    add_section_heading(doc, 'Prerequisites', 3)
    add_bullet(doc, '', 'Phase 3 exit criteria met for the specific site being cut over')
    add_bullet(doc, '', 'Approved maintenance window confirmed in writing')
    add_bullet(doc, '', 'DialConnection has coordinated communication to end customers')
    add_bullet(doc, '', 'Rollback plan documented and reviewed')
    add_section_heading(doc, 'Key Activities (per site)', 3)
    add_checklist_item(doc, 'Confirm Carbonite delta sync lag is minimal (<5 min) before initiating cutover window')
    add_checklist_item(doc, 'Active Directory domain controllers migrated first — confirm AD replication between sites before domain-joined server cutover')
    add_checklist_item(doc, 'Execute final Carbonite sync and quiesce replication for the site being cut over')
    add_checklist_item(doc, 'Thrive re-IPs servers at target Thrive environment where applicable')
    add_checklist_item(doc, 'DNS, routing, and public IP changes executed per approved cutover runbook')
    add_checklist_item(doc, 'SIP Gateway configuration updated to reflect new Thrive IPs — validate call routing')
    add_checklist_item(doc, 'DialConnection validates all applications and services post-cutover')
    add_checklist_item(doc, 'DialConnection coordinates communication to end customers during and after cutover')
    add_checklist_item(doc, 'Thrive issues Availability Notification for each completed location')
    add_checklist_item(doc, 'DialConnection reviews and confirms services within 5 business days of Availability Notification')
    add_section_heading(doc, 'Risks and Contingencies', 3)
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['SIP call routing fails after gateway re-IP at Chicago or Dallas', 'Medium', 'Critical', 'Test SIP registration and call routing in Thrive environment (Phase 3) before cutover. Rollback: re-activate source servers at TierPoint within cutover window. Document rollback decision point at T+1hr.'],
            ['Active Directory replication issues between Chicago and Dallas post-cutover', 'Low', 'High', 'Validate AD replication during Phase 3. Migrate Chicago and Dallas DCs before other domain-joined servers. Rollback if AD replication not healthy within 30 min.'],
            ['DNS TTL not reduced prior to cutover causing extended propagation delay', 'Medium', 'Medium', 'Reduce DNS TTLs to 300 seconds at least 24 hours before cutover window. Restore to normal TTL after cutover is confirmed stable.'],
            ['Carbonite final sync window exceeds maintenance window', 'Low', 'High', 'Monitor delta sync lag in Phase 3. If lag > 30 min consistently, delay cutover and investigate. Large SQL servers are highest risk — plan a dedicated cutover window for CSQLCST1.'],
            ['Customer-facing PBX interruption during Nashville/Atlanta cutover', 'Medium', 'Critical', 'PBX is multi-tenant and customer-facing — cutover window must be outside business hours. Validate NSH-DRGateway-PBX and NPBX01 at Thrive Atlanta before cutting DNS and SIP routing.'],
        ],
        col_widths=[1.9, 0.85, 0.7, 3.15]
    )
    add_section_heading(doc, 'Phase 4 Exit Criteria', 3)
    add_bullet(doc, '', 'All four sites fully cut over to ThriveCloud — no production workloads remaining at TierPoint')
    add_bullet(doc, '', 'Availability Notifications issued for all four sites')
    add_bullet(doc, '', 'DialConnection confirms all sites functioning as expected within 5 business days per site')

    # --- Phase 5 ---
    add_section_heading(doc, 'Phase 5 — Post-Cutover Validation and Handoff', 2)
    add_body(doc,
        'Phase 5 confirms full steady-state operation across all four Thrive locations, activates '
        'billing, applies the one-time credit, and formally transfers the account to the Thrive '
        'managed services team. DialConnection begins TierPoint decommission proceedings independently.'
    )
    add_section_heading(doc, 'Key Activities', 3)
    add_checklist_item(doc, 'All four sites fully operational on ThriveCloud — validate across 30 days post-cutover')
    add_checklist_item(doc, 'Backup and DRaaS replication health confirmed across all sites')
    add_checklist_item(doc, 'MDR monitoring and alerting validated — SOC receiving events from all log sources')
    add_checklist_item(doc, 'EPP agent health confirmed on all 90 servers — all reporting to Thrive SOC')
    add_checklist_item(doc, 'Server patching baseline confirmed — first patch cycle completed and reported')
    add_checklist_item(doc, 'RDS licensing validated for all 50 contracted users at Chicago and Dallas')
    add_checklist_item(doc, 'Annual DR failover test scheduled')
    add_checklist_item(doc, 'Billing activation confirmed for all sites per Service Order terms')
    add_checklist_item(doc, 'One-time credit of $207,165.60 applied to account')
    add_checklist_item(doc, 'Change Log submitted and reconciled for any scope adjustments (Carbonite count, VM counts, compute changes)')
    add_checklist_item(doc, 'Project formally closed — account transferred to Thrive steady-state managed services team')
    add_section_heading(doc, 'Risks and Contingencies', 3)
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['DialConnection delays formal site confirmation beyond 30-day billing trigger', 'Low', 'Medium', 'Billing activates 30 days after Availability Notification regardless of client confirmation per Service Order terms.'],
            ['Change Log reconciliation increases MRC due to additional VMs or compute', 'Medium', 'Low', 'MRC adjustments cannot fall below the original Service Order minimum per Change Log terms. DialConnection notified before adjustments are finalized.'],
            ['DR test failure discovered during Phase 5 annual test', 'Low', 'High', 'Thrive investigates and remediates replication issues. Annual test is required for DR SLA to remain in effect per Service Order.'],
        ],
        col_widths=[1.9, 0.85, 0.7, 3.15]
    )
    add_section_heading(doc, 'Phase 5 Exit Criteria', 3)
    add_bullet(doc, '', 'Billing activated for all four sites')
    add_bullet(doc, '', '$207,165.60 one-time credit applied')
    add_bullet(doc, '', 'DR test scheduled and Change Log reconciliation complete')
    add_bullet(doc, '', 'Project formally closed and account in steady-state managed services')

    # -----------------------------------------------------------------------
    # 13. Service Level Agreement Reference
    # -----------------------------------------------------------------------
    add_section_heading(doc, '13. Service Level Agreements', 1)
    add_table(doc,
        ['Classification', 'Priority', 'Severity', 'Initial Response', 'Method'],
        [
            ['Major Incident', 'P1', '1 — Emergency', '15 minutes', 'Phone call required to initiate'],
            ['Critical Incident', 'P2', '2 — Critical', '30 minutes', 'Incident ticket via email/portal'],
            ['Urgent Incident', 'P3', '3 — Urgent', '60 minutes', 'Incident ticket via email/portal'],
            ['Normal Incident', 'P4', '4 — Normal', '4 hours', 'Incident ticket via email/portal'],
            ['Request', 'P5', '5 — Low', '24 hours', 'Incident ticket via email/portal'],
        ],
        col_widths=[1.4, 0.6, 1.2, 1.2, 2.2]
    )
    add_sub_note(doc,
        'Major Incidents (P1) must be initiated via phone call. Portal, email, and voicemail submissions '
        'are not accepted at P1 until the call is received by a technical resource. Client devices must be '
        'powered on and responsive during scheduled service windows. All devices must be within vendor support lifecycle.'
    )
    add_body(doc, 'DRaaS SLA: Failover initiation within 15 minutes of confirmed disaster declaration. Recovery target: 40 VMs within 1 hour, then 10 VMs every 15 minutes. RPO < 15 minutes. Annual simulated recovery test required.')
    add_body(doc, 'Patching SLA: Critical/security OS patches deployed within 30 days of vendor release and Thrive approval. Credit trigger if >5% of contracted servers miss eligible patches in a calendar period.')

    # -----------------------------------------------------------------------
    # 14. Project Assumptions
    # -----------------------------------------------------------------------
    add_section_heading(doc, '14. Project Assumptions', 1)
    add_bullet(doc, 'Administrative Access: ', 'DialConnection provides full administrative access to all in-scope servers, hypervisor platforms, and management systems required for migration.')
    add_bullet(doc, 'Bandwidth: ', 'DialConnection ensures sufficient bandwidth between TierPoint and Thrive to support ongoing replication and meet DRaaS RPO requirements prior to cutover.')
    add_bullet(doc, 'Application Compatibility: ', 'DialConnection is responsible for confirming compatibility of all third-party applications. Thrive is not responsible for SQL or application-level issues arising from server migration.')
    add_bullet(doc, 'IP and Hostname Variables: ', 'DialConnection is responsible for identifying all environment-specific variables (IPs, hostnames, app configs) that require updates in the Thrive environment.')
    add_bullet(doc, 'Source Licensing: ', 'DialConnection provides or procures all Windows Server and SQL Server licensing at the TierPoint source. Thrive provides licensing at ThriveCloud as contracted.')
    add_bullet(doc, 'Decommission: ', 'Decommissioning TierPoint source servers post-migration is DialConnection\'s sole responsibility. Thrive is not involved in TierPoint contract termination.')
    add_bullet(doc, 'Maintenance Windows: ', 'DialConnection designates maintenance windows with appropriate business unit stakeholders. Impact to DialConnection customers during migration is DialConnection\'s responsibility to communicate.')
    add_bullet(doc, 'Scope Changes: ', 'All work outside this plan\'s stated scope will be assessed as a Change Order and may result in additional fees. Change Orders require written approval before execution.')
    add_bullet(doc, 'Windows XP VMs: ', 'Carbonite Migrate does not support Windows XP migrations. DC3VBVM01/02 (and unconfirmed DC3VBVM03/04) require a separate DialConnection decision: rebuild at Thrive or decommission.')

    # -----------------------------------------------------------------------
    # 15. Client Responsibilities
    # -----------------------------------------------------------------------
    add_section_heading(doc, '15. Client Responsibilities', 1)
    add_checklist_item(doc, 'Designate a Project POC available throughout implementation to approve requests, attend meetings, and communicate plans internally.')
    add_checklist_item(doc, 'Provide all credentials and administrative access to onboard all contracted Thrive products and services, including all source servers, hypervisors, and management platforms.')
    add_checklist_item(doc, 'Complete the DialConnection ThriveCloud Migration Information Gathering Workbook (all yellow fields) prior to Phase 1 close.')
    add_checklist_item(doc, 'Review and approve the Thrive-generated VM Specification list before provisioning begins.')
    add_checklist_item(doc, 'Provide DialConnection escalation contacts to Thrive for service event notification (MDR high-severity alerts, backup failures, DR events).')
    add_checklist_item(doc, 'Provide a Windows Domain Administrator account for automated patch deployment and remediation.')
    add_checklist_item(doc, 'Confirm compatibility of all third-party applications prior to deployment in the Thrive cloud platform.')
    add_checklist_item(doc, 'Designate appropriate maintenance/outage windows with relevant business units for each site cutover.')
    add_checklist_item(doc, 'Coordinate communication with DialConnection\'s end customers during the migration period.')
    add_checklist_item(doc, 'Validate and confirm application/service functionality at each site within 5 business days of each Availability Notification.')
    add_checklist_item(doc, 'Reboot servers when prompted during patching windows — ignoring reboot prompts hinders current and future patch deployments.')
    add_checklist_item(doc, 'Ensure production servers at TierPoint maintain sufficient CPU, memory, and storage resources to support continuous Carbonite replication.')
    add_checklist_item(doc, 'Provide full public IP usage mapping (per-IP service assignment) for Chicago and Dallas prior to Phase 4 cutover planning.')
    add_checklist_item(doc, 'Document all undocumented VLANs at Nashville (VLAN 3023 / PBX-Internal, VLAN 3128 / PROD-Public_2) prior to Atlanta cutover.')

    # -----------------------------------------------------------------------
    # 16. Open Items
    # -----------------------------------------------------------------------
    add_section_heading(doc, '16. Open Items', 1)
    add_body(doc, 'The following items require DialConnection input before the indicated phase can proceed.')
    add_table(doc,
        ['#', 'Item', 'Required By', 'Owner'],
        [
            ['1', 'All client contacts — Project POC, IT Lead, after-hours, site leads, app owners', 'Phase 1 kickoff', 'DialConnection'],
            ['2', 'Maintenance / cutover windows per site', 'Phase 3', 'DialConnection'],
            ['3', 'DNS configuration — domain names, DNS servers, zones, external DNS hosting', 'Phase 2', 'DialConnection'],
            ['4', 'Full public IP usage mapping — all assigned IPs at Chicago and Dallas', 'Phase 4', 'DialConnection'],
            ['5', 'VPN tunnel inventory — all IPSEC tunnels at Chicago (16) and Dallas (18)', 'Phase 2', 'DialConnection'],
            ['6', 'Remote access VPN details — platform, auth method, MFA provider, user count', 'Phase 2', 'DialConnection'],
            ['7', 'Firewall rule export from TierPoint FortiGates (all 4 sites)', 'Phase 2', 'DialConnection'],
            ['8', 'Migration wave assignments, target IPs, and DNS/FQDN per server', 'Phase 3', 'DialConnection'],
            ['9', 'DR recovery groups, boot order, and dependencies (partially pre-populated in workbook)', 'Phase 3', 'DialConnection'],
            ['10', 'DR test schedule', 'Phase 5', 'DialConnection'],
            ['11', 'Application dependency details — ports, tenant counts, database connections', 'Phase 3', 'DialConnection'],
            ['12', 'Confirm DC3VBVM03/04 existence — migrate, rebuild, or decommission decision', 'Phase 1', 'DialConnection'],
            ['13', 'Document Nashville VLAN 3023 (PBX-Internal 10.6.2.0/24) and VLAN 3128 (secondary public block)', 'Phase 2', 'DialConnection'],
            ['14', 'Confirm 3rd Linux VM in Boston (SO contracts 3 Linux; inventory shows 2)', 'Phase 1', 'DialConnection'],
            ['15', 'Carbonite license count reconciliation (90 contracted vs. 95 VM inventory) via Change Log', 'Phase 2', 'Thrive + DialConnection'],
            ['16', 'SQL Availability Group / Always On configuration documentation for Chicago and Dallas', 'Phase 3', 'DialConnection'],
            ['17', 'GatewayDALDR002/003/004 removal decision (flagged as possible decommission)', 'Phase 1', 'DialConnection'],
        ],
        col_widths=[0.3, 3.5, 1.3, 1.5]
    )

    # -----------------------------------------------------------------------
    # 17. Financial Summary
    # -----------------------------------------------------------------------
    add_section_heading(doc, '17. Financial Summary', 1)
    add_body(doc,
        'The following summarizes the contracted financial commitments under Service Order CON-094865. '
        'All amounts are in USD. Service Order is dated March 27, 2026.'
    )
    add_table(doc,
        ['Category', 'Service Order Line', 'Amount'],
        [
            ['One-Time — Implementation', 'Premium Implementation Services (fixed fee) + 4 ThriveCloud Migration Base Fees + 90 Block Level Migrations', '$200,685.74'],
            ['One-Time — Discount', 'Sales Provisioning Services One-Time Discount', '($138,000.00)'],
            ['One-Time Services Net', '', '$62,685.74'],
            ['MRC — ThriveCloud Atlanta', '', '$3,162.30'],
            ['MRC — ThriveCloud Boston', '', '$2,895.70'],
            ['MRC — ThriveCloud Chicago', '', '$54,979.30'],
            ['MRC — ThriveCloud Dallas', '', '$53,148.30'],
            ['MRC — Managed Security', 'MDR Essentials + Endpoint Security and Response', '$4,467.60'],
            ['MRC — Managed Patching', 'Enterprise Plus + Third-Party Advanced Patching (90 servers)', '$4,452.00'],
            ['MRC — Advisory Services', 'Technical Advisory Services (fully discounted)', '$0.00'],
            ['MRC — RDS Licensing', '50 Windows RDS Access Licenses (per user)', '$850.00'],
            ['MRC — Recurring Discount', 'Sales Managed Services Recurring Discount', '($54,900.00)'],
            ['Net Monthly Recurring Total', '', '$69,055.20'],
            ['Migration Software (one-time product)', '90 x Carbonite Migrate Standard for Windows and Linux', '$40,500.00'],
            ['One-Time Credit at Billing Activation', '~3 months MRC (non-transferable)', '($207,165.60)'],
        ],
        col_widths=[2.4, 2.5, 1.7]
    )
    add_sub_note(doc,
        'Contract term is 63 months. Billing activates per-site upon each Availability Notification. '
        'Early termination requires payment of all remaining fees through end of then-current term. '
        'The Change Log process per the Service Order allows for scope and quantity adjustments during '
        'onboarding without requiring a formal Change Order for each adjustment. MRC cannot fall below '
        'the contracted minimum per the Change Log terms.'
    )

    add_footer(doc)
    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')


if __name__ == '__main__':
    build_tip()
