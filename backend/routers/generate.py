"""
TIP generation API endpoints
"""
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
import json
import io
from database import get_db
from models.draft import Draft, DraftStatus, DraftCollaborator, DraftDocument
from models.document import Document
from models.template_file import TemplateFile
from models.user import User as UserModel
from schemas.draft import DraftCreate, DraftResponse, GenerateTIPRequest, GenerateTIPResponse, DraftUpdate, RefineRequest, RefineResponse
from services.claude import ClaudeService
from celery_app import generate_tip_task
from routers.auth import get_current_user

router = APIRouter(prefix="/api/generate", tags=["generate"])


def _is_collaborator(db: Session, draft_id: int, user_id: int) -> bool:
    return db.query(DraftCollaborator).filter(
        DraftCollaborator.draft_id == draft_id,
        DraftCollaborator.user_id == user_id,
    ).first() is not None


def _get_draft_readable(db: Session, draft_id: int, user: UserModel) -> Draft:
    """Return draft if user is owner, collaborator, or admin. 404 otherwise."""
    draft = db.query(Draft).filter(Draft.id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    if draft.user_id == user.id or user.is_superuser or _is_collaborator(db, draft_id, user.id):
        return draft
    raise HTTPException(status_code=404, detail="Draft not found")


def _get_draft_owned(db: Session, draft_id: int, user: UserModel) -> Draft:
    """Return draft only if user is owner or admin. 404 otherwise."""
    draft = db.query(Draft).filter(Draft.id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    if draft.user_id == user.id or user.is_superuser:
        return draft
    raise HTTPException(status_code=404, detail="Draft not found")


def get_claude_service_for_user(user: UserModel) -> ClaudeService:
    if not user.claude_api_key:
        from fastapi import HTTPException
        raise HTTPException(status_code=402, detail="No Claude API key configured. Add your Anthropic API key in your profile settings.")
    return ClaudeService(api_key=user.claude_api_key, model=user.claude_model or None)

@router.post("/draft", response_model=DraftResponse)
async def create_draft(
    draft_data: DraftCreate,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """
    Create a new TIP draft
    """
    # Validate documents exist if provided — allow user's own docs and shared docs from user ID 1
    from sqlalchemy import or_ as sql_or
    def _can_access_doc(doc_id: int) -> bool:
        return db.query(Document).filter(
            Document.id == doc_id,
            sql_or(Document.user_id == current_user.id, Document.user_id == 1)
        ).first() is not None

    if draft_data.discovery_document_id:
        if not _can_access_doc(draft_data.discovery_document_id):
            raise HTTPException(status_code=404, detail="Discovery document not found")
    
    if draft_data.service_order_document_id:
        if not _can_access_doc(draft_data.service_order_document_id):
            raise HTTPException(status_code=404, detail="Service order document not found")

    for sup_id in (draft_data.supplemental_document_ids or []):
        if not _can_access_doc(sup_id):
            raise HTTPException(status_code=404, detail=f"Supplemental document {sup_id} not found")

    # Create draft
    draft = Draft(
        user_id=current_user.id,
        title=draft_data.title,
        description=draft_data.description,
        discovery_document_id=draft_data.discovery_document_id,
        service_order_document_id=draft_data.service_order_document_id,
        template_id=draft_data.template_id,
        scrub_pii=draft_data.scrub_pii,
        status=DraftStatus.DRAFT
    )
    
    db.add(draft)
    db.commit()
    db.refresh(draft)

    # Populate draft_documents junction table
    pos = 0
    if draft_data.discovery_document_id:
        db.add(DraftDocument(draft_id=draft.id, document_id=draft_data.discovery_document_id, role="discovery", position=pos))
        pos += 1
    if draft_data.service_order_document_id:
        db.add(DraftDocument(draft_id=draft.id, document_id=draft_data.service_order_document_id, role="service_order", position=pos))
        pos += 1
    for sup_id in (draft_data.supplemental_document_ids or []):
        db.add(DraftDocument(draft_id=draft.id, document_id=sup_id, role="supplemental", position=pos))
        pos += 1
    db.commit()
    db.refresh(draft)

    return draft

@router.post("/tip", response_model=GenerateTIPResponse)
async def generate_tip(
    request: GenerateTIPRequest,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """
    Enqueue TIP generation as a Celery task and return immediately.
    Poll GET /drafts/{id} to check status (generating → completed/failed).
    """
    draft = db.query(Draft).filter(
        Draft.id == request.draft_id,
        Draft.user_id == current_user.id
    ).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    active_template = db.query(TemplateFile).filter(TemplateFile.is_active == True).first()
    template_file_id = active_template.id if active_template else None

    draft.status = DraftStatus.GENERATING
    db.commit()

    task = generate_tip_task.delay(draft.id, template_file_id)
    draft.celery_task_id = task.id
    db.commit()

    return GenerateTIPResponse(
        message="TIP generation started",
        draft_id=draft.id,
        status=DraftStatus.GENERATING,
        content=None
    )

@router.get("/drafts", response_model=List[DraftResponse])
async def list_drafts(
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """
    List all drafts owned by or shared with the current user.
    """
    from sqlalchemy import or_
    collab_draft_ids = db.query(DraftCollaborator.draft_id).filter(
        DraftCollaborator.user_id == current_user.id
    ).subquery()

    drafts = db.query(Draft)\
        .filter(or_(
            Draft.user_id == current_user.id,
            Draft.id.in_(collab_draft_ids),
        ))\
        .order_by(Draft.created_at.desc())\
        .offset(skip)\
        .limit(limit)\
        .all()

    return drafts

@router.patch("/drafts/{draft_id}", response_model=DraftResponse)
async def update_draft(
    draft_id: int,
    update: DraftUpdate,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    draft = _get_draft_readable(db, draft_id, current_user)
    draft.content = update.content
    if update.title:
        draft.title = update.title
    db.commit()
    db.refresh(draft)
    return draft

@router.post("/drafts/{draft_id}/refine", response_model=RefineResponse)
async def refine_draft(
    draft_id: int,
    request: RefineRequest,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    import asyncio
    from concurrent.futures import ThreadPoolExecutor
    draft = _get_draft_readable(db, draft_id, current_user)
    sections_content = "\n\n".join(draft.sections.values()) if draft.sections else ""
    if not draft.content and not sections_content and not request.current_content:
        raise HTTPException(status_code=400, detail="Draft has no content to refine")
    try:
        claude_service = get_claude_service_for_user(current_user)
        content = request.current_content or draft.content or sections_content or ""
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as pool:
            suggestion = await loop.run_in_executor(
                pool,
                lambda: claude_service.refine_tip(
                    instruction=request.instruction,
                    current_content=content
                )
            )
        return RefineResponse(suggestion=suggestion)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Refinement failed: {str(e)}")

@router.get("/drafts/{draft_id}", response_model=DraftResponse)
async def get_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """
    Get details of a specific draft
    """
    draft = _get_draft_readable(db, draft_id, current_user)
    return draft

@router.get("/drafts/{draft_id}/progress")
async def get_draft_progress(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """
    Lightweight polling endpoint — returns status and chunk progress only.
    Does NOT return content. Use this during generation instead of GET /drafts/{id}.
    """
    draft = _get_draft_readable(db, draft_id, current_user)

    progress = None
    if draft.generation_prompt:
        try:
            p = json.loads(draft.generation_prompt)
            if p.get("mode") == "chunked":
                progress = {
                    "chunk": p.get("chunk", 0),
                    "total_chunks": p.get("total_chunks", 0),
                    "sections": p.get("sections", 0),
                }
        except Exception:
            pass

    return {
        "draft_id": draft.id,
        "status": draft.status,
        "title": draft.title,
        "progress": progress,
        "generation_tokens": draft.generation_tokens,
    }


@router.post("/drafts/{draft_id}/cancel")
async def cancel_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Cancel an in-progress generation: revoke the Celery task and mark draft as failed."""
    from celery_app import celery
    draft = _get_draft_owned(db, draft_id, current_user)
    if draft.status != DraftStatus.GENERATING:
        raise HTTPException(status_code=400, detail="Draft is not currently generating")
    # Revoke task on all workers — stops it before next chunk boundary
    if draft.celery_task_id:
        try:
            celery.control.revoke(draft.celery_task_id, terminate=True, signal="SIGTERM", reply=False)
        except Exception:
            pass
    # Mark failed immediately so the running task self-terminates at next chunk check
    draft.status = DraftStatus.FAILED
    draft.content = "Generation cancelled by user."
    draft.celery_task_id = None
    draft.generation_prompt = None
    db.commit()
    return {"message": "Generation cancelled", "id": draft_id}


@router.delete("/drafts/{draft_id}")
async def delete_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    draft = _get_draft_owned(db, draft_id, current_user)
    db.delete(draft)
    db.commit()
    return {"message": "Draft deleted", "id": draft_id}


@router.post("/drafts/{draft_id}/duplicate", response_model=DraftResponse)
async def duplicate_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Create a copy of a draft owned by the current user."""
    original = _get_draft_readable(db, draft_id, current_user)
    copy = Draft(
        user_id=current_user.id,
        template_id=original.template_id,
        template_file_id=original.template_file_id,
        title=f"{original.title} (Copy)",
        description=original.description,
        status=DraftStatus.COMPLETED if original.status == DraftStatus.COMPLETED else DraftStatus.DRAFT,
        discovery_document_id=original.discovery_document_id,
        service_order_document_id=original.service_order_document_id,
        content=original.content,
        sections=dict(original.sections) if original.sections else None,
        claude_model=original.claude_model,
        generation_tokens=original.generation_tokens,
    )
    db.add(copy)
    db.commit()
    db.refresh(copy)
    return copy


@router.get("/drafts/{draft_id}/gaps")
async def get_draft_gaps(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Scan draft content for [DATA NEEDED: ...] placeholders and return them."""
    import re
    draft = _get_draft_readable(db, draft_id, current_user)

    text = draft.content or ""
    pattern = re.compile(r'\[DATA NEEDED:\s*(.*?)\]', re.IGNORECASE | re.DOTALL)

    gaps = []
    if draft.sections:
        for section_key, section_content in draft.sections.items():
            for m in pattern.finditer(section_content or ""):
                gaps.append({"section": section_key, "placeholder": m.group(0).strip(), "detail": m.group(1).strip()})
    else:
        for m in pattern.finditer(text):
            gaps.append({"section": None, "placeholder": m.group(0).strip(), "detail": m.group(1).strip()})

    return {"draft_id": draft_id, "gap_count": len(gaps), "gaps": gaps}


@router.patch("/drafts/{draft_id}/sections/{section_key:path}")
async def update_draft_section(
    draft_id: int,
    section_key: str,
    body: dict,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Update a single section of a draft by key.
    section_key may also be passed in body as 'key' to avoid URL slash issues.
    """
    draft = _get_draft_readable(db, draft_id, current_user)
    # Prefer key from body if present (avoids URL encoding issues with slashes)
    resolved_key = body.get("key") or section_key
    sections = dict(draft.sections or {})
    sections[resolved_key] = body.get("content", "")
    draft.sections = sections
    draft.content = "\n\n".join(
        f"## {k}\n\n{v}" for k, v in sections.items() if v
    )
    db.commit()
    db.refresh(draft)
    return {"section": resolved_key, "saved": True}


@router.get("/template-instructions")
async def get_template_instructions():
    """Return the parsed instruction map from the Thrive TIP template."""
    import os, json as _json
    path = os.path.join(os.path.dirname(__file__), '..', 'static', 'template_instructions.json')
    path = os.path.normpath(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Template instructions not yet extracted")
    with open(path) as f:
        return _json.load(f)


@router.post("/drafts/{draft_id}/refine-guided")
async def refine_section_guided(
    draft_id: int,
    body: dict,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Refine a single section using the template instruction for that section type."""
    import os, json as _json

    draft = _get_draft_readable(db, draft_id, current_user)

    import datetime as _dt
    from models.user import User as UserModel

    section_key = body.get("section_key", "")
    current_content = body.get("current_content", "")
    mode = body.get("mode", "tighten")  # tighten | comply | risks | both | custom
    custom_instruction = (body.get("custom_instruction") or "").strip()

    # Resolve author name from current user
    author_name = (current_user.full_name or current_user.username) if current_user else "Thrive"
    today = _dt.date.today().strftime("%B %d, %Y")

    # Load instruction map from active template in DB (overrides layer on top of parsed)
    from models.template_file import TemplateFile
    template_instructions: dict = {}
    active_tpl = db.query(TemplateFile).filter(TemplateFile.is_active == True).first()
    if active_tpl and active_tpl.template_structure:
        try:
            tpl_structure = _json.loads(active_tpl.template_structure)
            raw = tpl_structure.get("instructions", [])
            if isinstance(raw, list):
                for item in raw:
                    sec = (item.get("section") or "").strip()
                    txt = (item.get("text") or "").strip()
                    if sec and txt:
                        template_instructions[sec] = txt
            elif isinstance(raw, dict):
                template_instructions = raw
            # Overrides (manually edited in template admin) take precedence
            template_instructions.update(tpl_structure.get("instruction_overrides", {}))
        except Exception:
            pass
    # Fallback to static file if DB has nothing
    if not template_instructions:
        import os as _os
        instr_path = _os.path.normpath(_os.path.join(_os.path.dirname(__file__), '..', 'static', 'template_instructions.json'))
        if _os.path.exists(instr_path):
            with open(instr_path) as f:
                template_instructions = _json.load(f).get("instructions", {})

    # Hard-coded rules for sections the template doesn't have an explicit instruction block for
    HARD_RULES = {
        "revision history": (
            f"The Revision History table must use version number 1.0 for the initial release (never 0.1 or 0.9). "
            f"Format as a markdown table with columns: Rev #, Author(s), Change, Date. "
            f"The first and only row should be: | 1.0 | {author_name} | Initial Release | {today} |. "
            f"Today's date is {today}. The author is {author_name}. Do not use any other date or version number."
        ),
    }

    # Find the best matching instruction key — check hard rules first
    instruction_text = None
    for rule_key, rule_text in HARD_RULES.items():
        if rule_key in section_key.lower():
            instruction_text = rule_text
            break

    if not instruction_text:
        for key in template_instructions:
            if key and (key.lower() in section_key.lower() or section_key.lower() in key.lower()):
                instruction_text = template_instructions[key]
                break
    # Fallback: try substring match on H2-level keys
    if not instruction_text:
        h2_keys = ["Executive Summary", "Implementation Summary", "Requirements/Prerequisites",
                   "Approximate Timing", "Implementation Details", "Risks and Contingencies",
                   "Testing and Verification", "Day-1 Support", "Acceptance Criteria",
                   "Deliverables", "Timeline of Phases"]
        for key in h2_keys:
            if key.lower() in section_key.lower():
                instruction_text = template_instructions.get(key)
                break

    # Revision History: skip Claude entirely — stamp deterministically
    if "revision history" in section_key.lower():
        canonical = (
            "| Rev # | Author(s) | Change | Date |\n"
            "|-------|-----------|--------|------|\n"
            f"| 1.0 | {author_name} | Initial Release | {today} |"
        )
        return {
            "suggestion": canonical,
            "section_key": section_key,
            "instruction_used": f"Version 1.0, author = {author_name}, date = {today}",
            "mode": "direct"
        }

    mode_prompts = {
        "tighten": (
            "Tighten and condense this section. Remove all redundancy, wordiness, and over-explanation. "
            "Preserve every factual detail. Cut length by at least 40%. Be ruthlessly concise."
        ),
        "comply": (
            "Rewrite this section to precisely match the template instruction below. "
            "Remove anything that doesn't belong. Add any required structure. Keep all factual specifics."
        ),
        "risks": (
            "Convert this section into a 4-column markdown table matching the Thrive TIP template exactly. "
            "The table MUST have these exact headers: | Risk | Likelihood | Mitigation Strategy | Rollback Plan |\n"
            "Each row = one risk. Keep each cell to 1-2 concise sentences maximum. "
            "No prose paragraphs. No bullet points outside the table. Table rows only."
        ),
        "both": (
            "Tighten AND rewrite this section to match the template instruction. "
            "Remove all redundancy. Apply required structure. Preserve all factual details from the original."
        ),
    }

    if mode == "custom":
        if not custom_instruction:
            raise HTTPException(status_code=400, detail="Custom instruction is required for custom mode")
        mode_instruction = custom_instruction
    else:
        mode_instruction = mode_prompts.get(mode, mode_prompts["tighten"])

    instruction_block = f"\nTEMPLATE INSTRUCTION FOR THIS SECTION TYPE:\n{instruction_text}\n" if instruction_text else ""

    system_prompt = f"""You are a senior technical writer at Thrive Networks editing a Technical Implementation Plan (TIP).

{mode_instruction}
{instruction_block}
Rules:
- Return ONLY the revised section content, no preamble or explanation
- You MUST follow the template instruction above precisely — it is authoritative
- Preserve all customer-specific details, IP addresses, server names, dates
- Use markdown formatting (## headings, bullets, **bold** for key terms)
- Do not invent facts or add content not grounded in the original
- For Risks and Contingencies: output a markdown table with columns: Risk | Likelihood | Mitigation Strategy | Rollback Plan. One row per risk. Each cell = 1 sentence max. No prose.
- For Implementation Details: use numbered steps or clear sub-sections
- For Revision History: version 1.0 = initial release (never 0.1)
- Be concise. Do not pad responses. Less is more."""

    prompt = f"""Section: {section_key}

Current content:
{current_content}

Rewrite this section following the rules above."""

    if not current_user.claude_api_key:
        raise HTTPException(status_code=402, detail="No Claude API key configured. Add your Anthropic API key in your profile settings.")

    try:
        import anthropic
        import asyncio
        from concurrent.futures import ThreadPoolExecutor

        def _call_claude():
            client = anthropic.Anthropic(api_key=current_user.claude_api_key, default_headers={"X-Anthropic-Do-Not-Store": "true"})
            message = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=1500,
                system=system_prompt,
                messages=[{"role": "user", "content": prompt}]
            )
            return message.content[0].text

        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as pool:
            suggestion = await loop.run_in_executor(pool, _call_claude)

        return {
            "suggestion": suggestion,
            "section_key": section_key,
            "instruction_used": instruction_text or "general tighten",
            "mode": mode
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Refinement failed: {str(e)}")


@router.post("/drafts/{draft_id}/refine-all")
async def refine_all_sections(
    draft_id: int,
    body: dict,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Apply a free-text instruction to every section of the draft and return revised sections.

    Runs all section calls in parallel (ThreadPoolExecutor) with prompt caching on the
    system block — real-time response, no async batch latency.
    """
    import asyncio
    from concurrent.futures import ThreadPoolExecutor

    instruction = (body.get("instruction") or "").strip()
    if not instruction:
        raise HTTPException(status_code=400, detail="Instruction is required")

    if not current_user.claude_api_key:
        raise HTTPException(status_code=402, detail="No Claude API key configured. Add your Anthropic API key in your profile settings.")

    draft = _get_draft_readable(db, draft_id, current_user)
    if not draft.sections:
        raise HTTPException(status_code=400, detail="Draft has no sections to refine")

    # Skip cover-page and structural placeholder sections — they have no real content to refine
    SKIP_REFINE = {
        "technical implementation plan", "document end", "template usage guide",
        "revision history", "table of contents",
    }
    sections_to_refine = {
        k: v for k, v in draft.sections.items()
        if v and v.strip() and k.strip().lower() not in SKIP_REFINE
        and not k.strip().lower() == (draft.title or "").strip().lower()
    }
    if not sections_to_refine:
        raise HTTPException(status_code=400, detail="No refinable sections found")

    try:
        import anthropic as _anthropic

        REFINE_SYSTEM = [
            {
                "type": "text",
                "text": (
                    "You are a senior technical writer at Thrive Networks editing a Technical Implementation Plan (TIP). "
                    "Apply the user's instruction to the section content provided. "
                    "Return ONLY the revised section content — no preamble, no explanation, no section heading. "
                    "Preserve all customer-specific details, IP addresses, server names, dates. "
                    "Use markdown formatting. Do not invent facts."
                ),
                "cache_control": {"type": "ephemeral"},
            }
        ]

        model = current_user.claude_model or "claude-sonnet-4-5"

        def _refine_section(key: str, content: str) -> tuple[str, str]:
            client = _anthropic.Anthropic(api_key=current_user.claude_api_key, default_headers={"X-Anthropic-Do-Not-Store": "true"})
            prompt = f"Section: {key}\n\nInstruction: {instruction}\n\nCurrent content:\n{content[:8000]}\n\nApply the instruction above to this section."
            msg = client.messages.create(
                model=model,
                max_tokens=1500,
                system=REFINE_SYSTEM,
                messages=[{"role": "user", "content": prompt}]
            )
            return key, msg.content[0].text

        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor(max_workers=8) as pool:
            tasks = [
                loop.run_in_executor(pool, _refine_section, k, v)
                for k, v in sections_to_refine.items()
            ]
            results = await asyncio.gather(*tasks)

        return {"sections": dict(results), "instruction": instruction, "refined_count": len(results)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Refinement failed: {str(e)}")


@router.get("/drafts/{draft_id}/section-order")
async def get_section_order(draft_id: int, db: Session = Depends(get_db), current_user: UserModel = Depends(get_current_user)):
    """Return section order and visibility for a draft. Auto-initialises from draft.sections if not yet saved."""
    from models.draft_section_order import DraftSectionOrder
    draft = _get_draft_readable(db, draft_id, current_user)
    rows = db.query(DraftSectionOrder).filter(DraftSectionOrder.draft_id == draft_id).order_by(DraftSectionOrder.position).all()
    if not rows and draft.sections:
        # Bootstrap from current sections order
        rows = []
        for pos, key in enumerate(draft.sections.keys()):
            row = DraftSectionOrder(draft_id=draft_id, section_key=key, position=pos, visible=True)
            db.add(row)
            rows.append(row)
        db.commit()
    return [{"key": r.section_key, "position": r.position, "visible": r.visible} for r in rows]


@router.post("/drafts/{draft_id}/section-order")
async def save_section_order(draft_id: int, body: dict, db: Session = Depends(get_db), current_user: UserModel = Depends(get_current_user)):
    """
    Save section order and visibility for a draft.
    Body: { sections: [{ key: string, position: number, visible: boolean }] }
    """
    from models.draft_section_order import DraftSectionOrder
    _get_draft_readable(db, draft_id, current_user)
    sections = body.get("sections", [])
    for item in sections:
        key = item.get("key")
        if not key:
            continue
        row = db.query(DraftSectionOrder).filter(
            DraftSectionOrder.draft_id == draft_id,
            DraftSectionOrder.section_key == key
        ).first()
        if row:
            row.position = item.get("position", row.position)
            row.visible  = item.get("visible", row.visible)
        else:
            db.add(DraftSectionOrder(
                draft_id=draft_id,
                section_key=key,
                position=item.get("position", 0),
                visible=item.get("visible", True),
            ))
    db.commit()
    return {"saved": True, "count": len(sections)}


@router.get("/drafts/{draft_id}/export")
async def export_draft_docx(draft_id: int, db: Session = Depends(get_db), current_user: UserModel = Depends(get_current_user)):
    """Export a completed draft as a formatted Word (.docx) document."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.style import WD_STYLE_TYPE
    import re, copy

    draft = _get_draft_readable(db, draft_id, current_user)
    if not draft.content:
        raise HTTPException(status_code=400, detail="Draft has no content to export")

    import os

    # Resolve base template: prefer active DB template path, fallback to static file
    _active_tpl = db.query(TemplateFile).filter(TemplateFile.is_active == True).first()
    template_path = (
        _active_tpl.file_path
        if _active_tpl and _active_tpl.file_path and os.path.exists(_active_tpl.file_path)
        else os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'static', 'Thrive_TIP_Template.docx'))
    )
    if os.path.exists(template_path):
        doc = DocxDocument(template_path)
        # Clear all body content from template — keep styles/header/footer
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
        for tbl in list(doc.tables):
            t = tbl._element
            t.getparent().remove(t)
    else:
        doc = DocxDocument()

    # Header and footer are inherited directly from the template — do not modify them.

    def add_inline_runs(para, text: str, base_size=None):
        """Parse inline **bold**, *italic*, `code` and add runs to para."""
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                r = para.add_run(part[2:-2])
                r.font.bold = True
                if base_size: r.font.size = base_size
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                r = para.add_run(part[1:-1])
                r.font.italic = True
                if base_size: r.font.size = base_size
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                r = para.add_run(part[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
            else:
                r = para.add_run(part)
                if base_size: r.font.size = base_size

    def shade_cell(cell, hex_color: str):
        """Apply background fill to a single table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')):
            tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, color='CCCCCC', sz='4'):
        """Apply uniform border to all sides of a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def set_cell_padding(cell, top=80, bottom=80, left=120, right=120):
        """Set cell inner margin in twips."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(old)
        mar = OxmlElement('w:tcMar')
        for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            item = OxmlElement(f'w:{side}')
            item.set(qn('w:w'), str(val))
            item.set(qn('w:type'), 'dxa')
            mar.append(item)
        tcPr.append(mar)

    def set_cell_valign(cell, align='center'):
        """Set vertical alignment on a table cell."""
        from docx.enum.table import WD_ALIGN_VERTICAL
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if align == 'center' else WD_ALIGN_VERTICAL.TOP

    def set_para_spacing(para, before=0, after=60):
        """Set precise paragraph spacing via XML (twips)."""
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:spacing')):
            pPr.remove(old)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(before))
        spacing.set(qn('w:after'), str(after))
        pPr.append(spacing)

    # Pre-compute available style names once
    _available_styles = {s.name for s in doc.styles}

    def safe_style(name: str, fallback: str = 'Normal') -> str:
        return name if name in _available_styles else fallback

    def add_heading(text: str, level: int):
        """Add a heading using the template's built-in Heading styles."""
        style_name = safe_style(f'Heading {level}')
        p = doc.add_paragraph(text, style=style_name)
        if style_name == 'Normal':
            # Manual styling if Heading style not in template
            run = p.runs[0] if p.runs else p.add_run(text)
            sizes = {1: Pt(14), 2: Pt(12), 3: Pt(11)}
            run.font.size = sizes.get(level, Pt(11))
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x14, 0x3F, 0x6A)
        return p

    # Numbering state — tracked across list items
    _numbering_el = None
    try:
        _numbering_el = doc.part.numbering_part._element
    except Exception:
        pass

    # abstractNum IDs created once per type (reused by all num entries)
    _abstract_ids = {}

    def _ensure_abstract_num(is_ordered: bool) -> str:
        """Create abstractNum definition once per list type. Returns abstractNumId."""
        key = 'ordered' if is_ordered else 'bullet'
        if key in _abstract_ids:
            return _abstract_ids[key]
        if _numbering_el is None:
            return None
        from docx.oxml import OxmlElement as OE
        existing = _numbering_el.findall(qn('w:abstractNum'))
        abstract_id = str(len(existing))
        an = OE('w:abstractNum')
        an.set(qn('w:abstractNumId'), abstract_id)
        mlt = OE('w:multiLevelType'); mlt.set(qn('w:val'), 'hybridMultilevel'); an.append(mlt)
        lvl = OE('w:lvl'); lvl.set(qn('w:ilvl'), '0')
        start = OE('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
        if is_ordered:
            nf = OE('w:numFmt'); nf.set(qn('w:val'), 'decimal'); lvl.append(nf)
            lt = OE('w:lvlText'); lt.set(qn('w:val'), '%1.'); lvl.append(lt)
        else:
            nf = OE('w:numFmt'); nf.set(qn('w:val'), 'bullet'); lvl.append(nf)
            lt = OE('w:lvlText'); lt.set(qn('w:val'), '\u2022'); lvl.append(lt)
        lj = OE('w:lvlJc'); lj.set(qn('w:val'), 'left'); lvl.append(lj)
        pp = OE('w:pPr'); ind = OE('w:ind')
        ind.set(qn('w:left'), '720'); ind.set(qn('w:hanging'), '360')
        pp.append(ind); lvl.append(pp)
        rp = OE('w:rPr'); rf = OE('w:rFonts')
        rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
        rp.append(rf); lvl.append(rp)
        an.append(lvl)
        _numbering_el.append(an)
        _abstract_ids[key] = abstract_id
        return abstract_id

    def _new_num_id(is_ordered: bool) -> str:
        """Create a new w:num entry (= new list instance, counter resets to 1)."""
        abstract_id = _ensure_abstract_num(is_ordered)
        if abstract_id is None or _numbering_el is None:
            return None
        from docx.oxml import OxmlElement as OE
        existing = _numbering_el.findall(qn('w:num'))
        num_id = str(len(existing) + 1)
        num_el = OE('w:num'); num_el.set(qn('w:numId'), num_id)
        ref = OE('w:abstractNumId'); ref.set(qn('w:val'), abstract_id)
        num_el.append(ref)
        _numbering_el.append(num_el)
        return num_id

    # Track list continuity: reuse numId while in same list, reset when broken
    _cur_ordered_id = [None]   # mutable via closure
    _cur_bullet_id = [None]
    _last_was_ordered = [False]
    _last_was_bullet = [False]

    def _get_list_num_id(is_ordered: bool) -> str:
        """Return numId for current list group; create new one if list was interrupted."""
        if is_ordered:
            if not _last_was_ordered[0] or _cur_ordered_id[0] is None:
                _cur_ordered_id[0] = _new_num_id(True)
            _last_was_ordered[0] = True
            _last_was_bullet[0] = False
            return _cur_ordered_id[0]
        else:
            if _cur_bullet_id[0] is None:
                _cur_bullet_id[0] = _new_num_id(False)
            _last_was_bullet[0] = True
            _last_was_ordered[0] = False
            return _cur_bullet_id[0]

    def _break_list():
        """Call when a non-list element is rendered — signals list interruption."""
        _last_was_ordered[0] = False
        _last_was_bullet[0] = False

    def add_list_paragraph(text: str, is_ordered: bool = False, level: int = 0):
        """Add a properly formatted bullet or numbered list paragraph."""
        p = doc.add_paragraph(style='List Paragraph')
        set_para_spacing(p, before=0, after=40)
        num_id = _get_list_num_id(is_ordered)
        if num_id is not None:
            pPr = p._p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            ilvl_el = OxmlElement('w:ilvl')
            ilvl_el.set(qn('w:val'), str(level))
            numId_el = OxmlElement('w:numId')
            numId_el.set(qn('w:val'), num_id)
            numPr.append(ilvl_el)
            numPr.append(numId_el)
            pPr.append(numPr)
        else:
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            prefix = f'{text[:2]}' if is_ordered else '\u2022  '
            text = prefix + text
        add_inline_runs(p, text, base_size=Pt(10.5))
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        return p

    def add_horizontal_rule(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '143F6A')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def flush_table(tbl_lines):
        rows = [l for l in tbl_lines if not re.match(r'^\s*\|[-| :]+\|\s*$', l)]
        if not rows:
            return
        cols = max(len(r.strip('|').split('|')) for r in rows)
        t = doc.add_table(rows=0, cols=max(cols, 1))
        t.style = 'Table Grid'
        for ri, row_line in enumerate(rows):
            cells = [c.strip() for c in row_line.strip('|').split('|')]
            tr = t.add_row()
            if ri == 0:
                tr.height = Inches(0.33)
            for ci in range(max(cols, 1)):
                cell_text = cells[ci] if ci < len(cells) else ''
                cell_obj = tr.cells[ci]
                cell_obj.text = ''
                set_cell_valign(cell_obj)
                set_cell_padding(cell_obj, top=60, bottom=60, left=120, right=120)
                cell_para = cell_obj.paragraphs[0]
                set_para_spacing(cell_para, before=0, after=0)
                if ri == 0:
                    shade_cell(cell_obj, '143F6A')
                    set_cell_border(cell_obj, color='143F6A', sz='4')
                    add_inline_runs(cell_para, cell_text, base_size=Pt(10))
                    for r in cell_para.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.name = 'Calibri'
                else:
                    bg = 'F7F7F7' if ri % 2 == 1 else 'FFFFFF'
                    shade_cell(cell_obj, bg)
                    set_cell_border(cell_obj, color='CCCCCC', sz='4')
                    set_cell_padding(cell_obj, top=40, bottom=40, left=100, right=100)
                    add_inline_runs(cell_para, cell_text, base_size=Pt(10))
                    for r in cell_para.runs:
                        r.font.name = 'Calibri'
                        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Apply section order / visibility if saved — rebuild content from ordered visible sections
    if draft.sections:
        try:
            from models.draft_section_order import DraftSectionOrder
            order_rows = db.query(DraftSectionOrder).filter(
                DraftSectionOrder.draft_id == draft.id
            ).order_by(DraftSectionOrder.position).all()
            if order_rows:
                visible_keys = [r.section_key for r in order_rows if r.visible]
                ordered_content_parts = []
                for key in visible_keys:
                    val = draft.sections.get(key, "")
                    if val and val.strip():
                        level = 1 if not any(c.isalpha() and c.islower() for c in key[:3]) else 2
                        # Use heading level from content if detectable
                        ordered_content_parts.append(f"## {key}\n\n{val}")
                if ordered_content_parts:
                    content = "\n\n".join(ordered_content_parts)
                else:
                    content = draft.content or ""
            else:
                content = draft.content or ""
        except Exception:
            content = draft.content or ""
    else:
        content = draft.content or ""

    # Strip [INSTRUCTION: ...] blocks (may span multiple lines)
    content = re.sub(r'\[INSTRUCTION:.*?\]', '', content, flags=re.DOTALL)
    # Strip cover-page H1 ("# Technical Implementation Plan") and its immediate H2 subtitle
    content = re.sub(r'^# Technical Implementation Plan\s*\n', '', content, flags=re.MULTILINE)
    content = re.sub(r'^## .+ — .+Technical Implementation Plan\s*\n', '', content, flags=re.MULTILINE)
    # Strip Document Control Notice block (> blockquote lines)
    content = re.sub(r'(?:^> .*\n)+', '', content, flags=re.MULTILINE)
    # Strip standalone DOCUMENT CONTROL NOTICE paragraph
    content = re.sub(
        r'^\*{0,2}DOCUMENT CONTROL NOTICE\*{0,2}\n(?:.*\n)*?(?=\n#{1,3}|\n\n#{1,3})',
        '', content, flags=re.MULTILINE
    )
    # Strip Service Order callout lines (bold header + value lines at top of doc)
    content = re.sub(
        r'^\*\*Service Order:\*\*.*$', '', content, flags=re.MULTILINE
    )
    content = re.sub(
        r'^\*\*Prepared by:\*\*.*$', '', content, flags=re.MULTILINE
    )
    content = re.sub(
        r'^\*\*Date:\*\*.*$', '', content, flags=re.MULTILINE
    )
    # Collapse runs of 3+ blank lines to 2
    content = re.sub(r'\n{3,}', '\n\n', content)

    lines = content.split('\n')
    i = 0
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Table detection
        if line.strip().startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            flush_table(table_lines)
            table_lines = []
            _break_list()
            # fall through to process current line

        # Skip instruction remnants
        if re.match(r'^\[INSTRUCTION:', line.strip()):
            i += 1
            continue

        # Blockquote / callout
        if line.startswith('> '):
            _break_list()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:color'), '143F6A')
            pBdr.append(left)
            pPr.append(pBdr)
            add_inline_runs(p, line[2:].strip())
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x14, 0x3F, 0x6A)
                r.font.italic = True
            i += 1
            continue

        # Headings — use proper Word Heading styles from template
        if line.startswith('### '):
            _break_list()
            add_heading(line[4:].strip(), 3)
            i += 1
            continue
        if line.startswith('## '):
            _break_list()
            add_heading(line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith('# '):
            _break_list()
            add_heading(line[2:].strip(), 1)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            _break_list()
            add_horizontal_rule(doc)
            i += 1
            continue

        # Bullet / checklist
        if line.startswith('- [ ] ') or line.startswith('[ ] '):
            add_list_paragraph(f'\u2610  {line.lstrip("- ").lstrip("[ ] ").strip()}', is_ordered=False)
            i += 1
            continue
        if line.startswith('- [x] ') or line.startswith('[x] '):
            add_list_paragraph(f'\u2611  {line.lstrip("- ").lstrip("[x] ").strip()}', is_ordered=False)
            i += 1
            continue
        if line.startswith('- ') or line.startswith('* '):
            add_list_paragraph(line[2:], is_ordered=False)
            i += 1
            continue

        # Numbered list  (1. text)
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if num_match:
            add_list_paragraph(num_match.group(2), is_ordered=True)
            i += 1
            continue

        # Aligned paragraph passthrough — <p style="text-align:center|right|justify">
        align_match = re.match(
            r'^<p\s+style=["\']text-align:(center|right|justify)["\']>(.*?)</p>$',
            line.strip(), re.IGNORECASE
        )
        if align_match:
            _break_list()
            align_val = align_match.group(1).lower()
            inner = align_match.group(2)
            p = doc.add_paragraph()
            set_para_spacing(p, before=0, after=80)
            add_inline_runs(p, inner, base_size=Pt(10.5))
            alignment_map = {
                'center':  WD_ALIGN_PARAGRAPH.CENTER,
                'right':   WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            p.paragraph_format.alignment = alignment_map.get(align_val, WD_ALIGN_PARAGRAPH.LEFT)
            for run in p.runs:
                run.font.name = 'Calibri'
                if not run.font.color.type:
                    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            i += 1
            continue

        # Normal paragraph — 10.5pt Calibri #222 with precise spacing
        _break_list()
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=80)
        add_inline_runs(p, line, base_size=Pt(10.5))
        for run in p.runs:
            run.font.name = 'Calibri'
            if not run.font.color.type:
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        i += 1

    # Flush any trailing table
    if table_lines:
        flush_table(table_lines)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r'[^\w\s-]', '', draft.title).strip().replace(' ', '_')
    filename = f"{safe_title}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/drafts/{draft_id}/export/pdf")
async def export_draft_pdf(draft_id: int, db: Session = Depends(get_db), current_user: UserModel = Depends(get_current_user)):
    """Export a completed draft as a PDF using LibreOffice headless conversion."""
    import subprocess
    import tempfile
    import os as _os
    import re as _re

    draft = _get_draft_readable(db, draft_id, current_user)
    if not draft.content:
        raise HTTPException(status_code=400, detail="Draft has no content to export")

    # Build docx bytes by calling the existing export route handler directly
    response = await export_draft_docx(draft_id, db, current_user)
    docx_bytes = b"".join([chunk async for chunk in response.body_iterator])

    safe_title = _re.sub(r'[^\w\s-]', '', draft.title).strip().replace(' ', '_')

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = _os.path.join(tmpdir, f"{safe_title}.docx")
        pdf_path = _os.path.join(tmpdir, f"{safe_title}.pdf")

        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        # LibreOffice headless conversion
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, docx_path
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0 or not _os.path.exists(pdf_path):
            raise HTTPException(
                status_code=500,
                detail=f"PDF conversion failed: {result.stderr or result.stdout}"
            )

        pdf_bytes = open(pdf_path, 'rb').read()

    filename = f"{safe_title}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ---------------------------------------------------------------------------
# Collaborator endpoints
# ---------------------------------------------------------------------------

class CollaboratorResponse(BaseModel):
    user_id: int
    username: str
    full_name: Optional[str]
    invited_by_username: str
    created_at: datetime

    class Config:
        from_attributes = True


class InviteCollaboratorRequest(BaseModel):
    username: str  # invite by username


@router.get("/drafts/{draft_id}/collaborators", response_model=List[CollaboratorResponse])
async def list_collaborators(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """List all collaborators on a draft. Accessible by owner, collaborators, and admins."""
    _get_draft_readable(db, draft_id, current_user)
    rows = db.query(DraftCollaborator).filter(DraftCollaborator.draft_id == draft_id).all()
    result = []
    for row in rows:
        inviter = db.query(UserModel).filter(UserModel.id == row.invited_by).first()
        result.append(CollaboratorResponse(
            user_id=row.user_id,
            username=row.user.username,
            full_name=row.user.full_name,
            invited_by_username=inviter.username if inviter else "unknown",
            created_at=row.created_at,
        ))
    return result


@router.post("/drafts/{draft_id}/collaborators", response_model=CollaboratorResponse)
async def add_collaborator(
    draft_id: int,
    body: InviteCollaboratorRequest,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Invite a user to collaborate on a draft. Owner and admins only."""
    draft = _get_draft_owned(db, draft_id, current_user)

    invitee = db.query(UserModel).filter(UserModel.username == body.username).first()
    if not invitee:
        raise HTTPException(status_code=404, detail=f"User '{body.username}' not found")
    if invitee.id == draft.user_id:
        raise HTTPException(status_code=400, detail="Cannot invite the draft owner as a collaborator")

    existing = db.query(DraftCollaborator).filter(
        DraftCollaborator.draft_id == draft_id,
        DraftCollaborator.user_id == invitee.id,
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail=f"User '{body.username}' is already a collaborator")

    collab = DraftCollaborator(
        draft_id=draft_id,
        user_id=invitee.id,
        invited_by=current_user.id,
    )
    db.add(collab)
    db.commit()
    db.refresh(collab)

    return CollaboratorResponse(
        user_id=invitee.id,
        username=invitee.username,
        full_name=invitee.full_name,
        invited_by_username=current_user.username,
        created_at=collab.created_at,
    )


@router.delete("/drafts/{draft_id}/collaborators/{user_id}")
async def remove_collaborator(
    draft_id: int,
    user_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user),
):
    """Remove a collaborator from a draft. Owner, admins, or the collaborator themselves."""
    draft = db.query(Draft).filter(Draft.id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    # Allow owner, admin, or the collaborator removing themselves
    if draft.user_id != current_user.id and not current_user.is_superuser and current_user.id != user_id:
        raise HTTPException(status_code=404, detail="Draft not found")

    row = db.query(DraftCollaborator).filter(
        DraftCollaborator.draft_id == draft_id,
        DraftCollaborator.user_id == user_id,
    ).first()
    if not row:
        raise HTTPException(status_code=404, detail="Collaborator not found")

    db.delete(row)
    db.commit()
    return {"message": "Collaborator removed", "draft_id": draft_id, "user_id": user_id}
