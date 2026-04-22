"""
Library router — admin-managed reference TIP documents, globally visible to all users

Endpoints:
  GET    /api/library                   → list approved library docs (all users)
  GET    /api/library/all               → list all library docs including pending (admin)
  POST   /api/library                   → upload new library doc (admin)
  PATCH  /api/library/{id}/approve      → approve a pending doc (admin)
  PATCH  /api/library/{id}/reject       → reject a pending doc (admin)
  DELETE /api/library/{id}              → delete a library doc (admin)
  GET    /api/library/categories        → list distinct categories in use
"""
import os
import uuid
import shutil
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy.orm import Session

from database import get_db
from models.library import LibraryDocument, LibraryStatus
from models.user import User
from routers.auth import get_current_user
from routers.admin import require_admin
from services.upload import UploadService

router = APIRouter(prefix="/api/library", tags=["library"], redirect_slashes=False)

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/opt/tip-generator/uploads")
LIBRARY_DIR = os.path.join(UPLOAD_DIR, "library")

ALLOWED_EXTENSIONS = {".docx", ".pdf"}
MAX_SIZE_BYTES = 50 * 1024 * 1024  # 50MB


# --- Schemas ---

class LibraryDocumentResponse(BaseModel):
    id: int
    title: str
    category: str
    category_suggested: bool
    description: Optional[str]
    original_filename: str
    file_size: Optional[int]
    mime_type: Optional[str]
    status: str
    uploaded_by_username: str
    approved_by_username: Optional[str]
    approved_at: Optional[datetime]
    created_at: datetime

    class Config:
        from_attributes = True


class LibraryApprovalResponse(BaseModel):
    id: int
    status: str
    approved_at: Optional[datetime]


class LibraryCategoryUpdate(BaseModel):
    category: str


# --- Helpers ---

def _build_response(doc: LibraryDocument, db: Session) -> LibraryDocumentResponse:
    uploader = db.query(User).filter(User.id == doc.uploaded_by).first()
    approver = db.query(User).filter(User.id == doc.approved_by).first() if doc.approved_by else None
    return LibraryDocumentResponse(
        id=doc.id,
        title=doc.title,
        category=doc.category or "",
        category_suggested=bool(doc.category_suggested),
        description=doc.description,
        original_filename=doc.original_filename,
        file_size=doc.file_size,
        mime_type=doc.mime_type,
        status=doc.status.value,
        uploaded_by_username=uploader.username if uploader else "unknown",
        approved_by_username=approver.username if approver else None,
        approved_at=doc.approved_at,
        created_at=doc.created_at,
    )


def _suggest_category(
    title: str,
    filename: str,
    extracted_text: Optional[str],
    api_key: Optional[str],
    model: str,
) -> Optional[str]:
    """Call Claude to suggest a category based on document title, filename, and a text preview."""
    if not api_key:
        return None
    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=api_key)
        preview = (extracted_text or "")[:3000]
        prompt = (
            "You are categorizing a Technical Implementation Plan (TIP) document for a Managed Service Provider library.\n"
            f"Document title: {title}\n"
            f"Filename: {filename}\n"
            f"Content preview:\n{preview}\n\n"
            "Reply with ONLY a short category label (2-5 words) that best describes this TIP type. "
            "Examples: 'M365 Migration', 'Azure Infrastructure', 'Network Deployment', "
            "'Server Consolidation', 'Security Hardening', 'VoIP Implementation'. "
            "No explanation, no punctuation — just the category label."
        )
        response = client.messages.create(
            model=model,
            max_tokens=20,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.content[0].text.strip().strip('"').strip("'")
    except Exception as e:
        print(f"[library] category suggestion failed: {e}")
        return None


def _extract_text(file_path: str, mime_type: str) -> Optional[str]:
    """Best-effort text extraction for library docs (Word/PDF)."""
    try:
        if mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif mime_type == "application/pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        print(f"Warning: text extraction failed for {file_path}: {e}")
    return None


# --- Routes ---

@router.get("/categories", response_model=List[str])
def list_categories(
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    """Return distinct categories from approved library documents."""
    rows = (
        db.query(LibraryDocument.category)
        .filter(LibraryDocument.status == LibraryStatus.APPROVED)
        .distinct()
        .order_by(LibraryDocument.category)
        .all()
    )
    return [r[0] for r in rows]


@router.get("", response_model=List[LibraryDocumentResponse])
def list_approved(
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    """List all approved library documents — visible to all authenticated users."""
    docs = (
        db.query(LibraryDocument)
        .filter(LibraryDocument.status == LibraryStatus.APPROVED)
        .order_by(LibraryDocument.category, LibraryDocument.title)
        .all()
    )
    return [_build_response(d, db) for d in docs]


@router.get("/all", response_model=List[LibraryDocumentResponse])
def list_all(
    db: Session = Depends(get_db),
    _admin: User = Depends(require_admin),
):
    """List all library documents including pending/rejected — admin only."""
    docs = (
        db.query(LibraryDocument)
        .order_by(LibraryDocument.status, LibraryDocument.created_at.desc())
        .all()
    )
    return [_build_response(d, db) for d in docs]


@router.post("", response_model=LibraryDocumentResponse)
async def upload_library_doc(
    file: UploadFile = File(...),
    title: str = Form(...),
    category: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    admin: User = Depends(require_admin),
):
    """Upload a new library reference TIP — admin only. Category is optional; Claude will suggest one if omitted."""
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Allowed types: {', '.join(ALLOWED_EXTENSIONS)}")

    content = await file.read()
    if len(content) > MAX_SIZE_BYTES:
        raise HTTPException(status_code=400, detail="File too large. Maximum 50MB.")

    os.makedirs(LIBRARY_DIR, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex}{ext}"
    file_path = os.path.join(LIBRARY_DIR, safe_name)

    with open(file_path, "wb") as f:
        f.write(content)

    mime_map = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
    }
    mime_type = mime_map.get(ext, "application/octet-stream")
    extracted = _extract_text(file_path, mime_type)

    # Resolve category — use provided value or ask Claude
    category_str = (category or "").strip()
    category_suggested = False
    if not category_str:
        claude_model = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-5")
        suggested = _suggest_category(
            title=title.strip(),
            filename=file.filename,
            extracted_text=extracted,
            api_key=admin.claude_api_key,
            model=claude_model,
        )
        if suggested:
            category_str = suggested
            category_suggested = True

    doc = LibraryDocument(
        title=title.strip(),
        category=category_str,
        category_suggested=category_suggested,
        description=description.strip() if description else None,
        filename=safe_name,
        original_filename=file.filename,
        file_path=file_path,
        file_size=len(content),
        mime_type=mime_type,
        status=LibraryStatus.PENDING,
        extracted_text=extracted,
        uploaded_by=admin.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return _build_response(doc, db)


@router.patch("/{doc_id}/set-category")
def update_category(
    doc_id: int,
    body: LibraryCategoryUpdate,
    db: Session = Depends(get_db),
    _admin: User = Depends(require_admin),
):
    """Update the category of a library document — clears the suggested flag."""
    doc = db.query(LibraryDocument).filter(LibraryDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Library document not found")
    doc.category = body.category.strip()
    doc.category_suggested = False
    db.commit()
    return _build_response(doc, db)


@router.patch("/{doc_id}/approve", response_model=LibraryApprovalResponse)
def approve_doc(
    doc_id: int,
    db: Session = Depends(get_db),
    admin: User = Depends(require_admin),
):
    """Approve a pending library document — admin only."""
    doc = db.query(LibraryDocument).filter(LibraryDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Library document not found")
    doc.status = LibraryStatus.APPROVED
    doc.approved_by = admin.id
    doc.approved_at = datetime.now(timezone.utc)
    db.commit()
    return LibraryApprovalResponse(id=doc.id, status=doc.status.value, approved_at=doc.approved_at)


@router.patch("/{doc_id}/reject", response_model=LibraryApprovalResponse)
def reject_doc(
    doc_id: int,
    db: Session = Depends(get_db),
    admin: User = Depends(require_admin),
):
    """Reject a pending library document — admin only."""
    doc = db.query(LibraryDocument).filter(LibraryDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Library document not found")
    doc.status = LibraryStatus.REJECTED
    db.commit()
    return LibraryApprovalResponse(id=doc.id, status=doc.status.value, approved_at=doc.approved_at)


@router.delete("/{doc_id}")
def delete_doc(
    doc_id: int,
    db: Session = Depends(get_db),
    _admin: User = Depends(require_admin),
):
    """Delete a library document and its file — admin only."""
    doc = db.query(LibraryDocument).filter(LibraryDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Library document not found")

    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception as e:
            print(f"Warning: could not delete file {doc.file_path}: {e}")

    db.delete(doc)
    db.commit()
    return {"message": "Library document deleted", "id": doc_id}
