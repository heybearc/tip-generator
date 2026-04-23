"""
File upload service
"""
import os
import uuid
from pathlib import Path
from typing import Optional
from fastapi import UploadFile
from sqlalchemy.orm import Session
from models.document import Document, DocumentType, DocumentStatus
import openpyxl
import fitz  # PyMuPDF
from docx import Document as DocxDocument

class UploadService:
    """Service for handling file uploads and processing"""
    
    def __init__(self, upload_dir: str = "/mnt/tip-uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
    async def save_upload(
        self,
        file: UploadFile,
        document_type: DocumentType,
        user_id: int,
        db: Session
    ) -> Document:
        """
        Save uploaded file and create database record
        """
        # Generate unique filename
        file_ext = Path(file.filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = self.upload_dir / unique_filename
        
        # Save file to disk
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Create database record
        document = Document(
            user_id=user_id,
            filename=unique_filename,
            original_filename=file.filename,
            file_path=str(file_path),
            file_size=len(content),
            mime_type=file.content_type,
            document_type=document_type,
            status=DocumentStatus.UPLOADING
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Process file in background (for now, just mark as completed)
        try:
            extracted_text = await self.extract_text(file_path, file_ext)
            document.extracted_text = extracted_text
            document.status = DocumentStatus.COMPLETED
        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)
        
        db.commit()
        db.refresh(document)
        
        return document
    
    async def extract_text(self, file_path: Path, file_ext: str) -> Optional[str]:
        """
        Extract text content from uploaded file
        """
        try:
            if file_ext.lower() in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif file_ext.lower() == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext.lower() == '.docx':
                return self._extract_from_docx(file_path)
            else:
                return None
        except Exception as e:
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_from_excel(self, file_path: Path) -> str:
        """
        Extract text from Excel discovery worksheet.
        Handles multi-sheet workbooks with key-value pairs, tables, and section headers.
        Produces structured text that Claude can use to populate TIP sections.
        """
        # Load twice: data_only for formula results, normal for merged cell structure
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
        except Exception:
            wb = openpyxl.load_workbook(file_path)

        # Build a merged-cell value map: each cell in a merged range gets the top-left value
        def build_merge_map(sheet):
            merge_map = {}
            for merged_range in sheet.merged_cells.ranges:
                top_left = sheet.cell(merged_range.min_row, merged_range.min_col)
                val = top_left.value
                for row in range(merged_range.min_row, merged_range.max_row + 1):
                    for col in range(merged_range.min_col, merged_range.max_col + 1):
                        if not (row == merged_range.min_row and col == merged_range.min_col):
                            merge_map[(row, col)] = val
            return merge_map

        output_parts = []

        # Build dropdown lookup: {(row, col): [allowed values]} from data validation
        def build_dropdown_map(sheet):
            dv_map = {}
            for dv in sheet.data_validations.dataValidation:
                if dv.type == "list" and dv.formula1:
                    raw = dv.formula1.strip('"')
                    options = [v.strip() for v in raw.split(",") if v.strip()]
                    if options:
                        for cell_range in str(dv.sqref).split():
                            try:
                                from openpyxl.utils import range_boundaries
                                min_col, min_row, max_col, max_row = range_boundaries(cell_range)
                                for r in range(min_row, max_row + 1):
                                    for c in range(min_col, max_col + 1):
                                        dv_map[(r, c)] = options
                            except Exception:
                                pass
            return dv_map

        for sheet_name in wb.sheetnames:
            try:
                sheet = wb[sheet_name]
                merge_map = build_merge_map(sheet)
                dropdown_map = build_dropdown_map(sheet)

                # Collect non-empty rows, resolving merged cells
                rows = []
                for row_idx, row in enumerate(sheet.iter_rows(), start=1):
                    cells = []
                    for col_idx, cell in enumerate(row, start=1):
                        val = merge_map.get((row_idx, col_idx), cell.value)
                        # Append dropdown hint if cell has validation options and no value
                        if val is None and (row_idx, col_idx) in dropdown_map:
                            opts = dropdown_map[(row_idx, col_idx)]
                            val = f"[Options: {', '.join(opts[:8])}]"
                        cells.append(str(val).strip() if val is not None else "")
                    # Strip trailing empty cells
                    while cells and cells[-1] == "":
                        cells.pop()
                    if any(cells):
                        rows.append(cells)

                if not rows:
                    continue

                output_parts.append(f"\n{'='*60}")
                output_parts.append(f"SECTION: {sheet_name}")
                output_parts.append(f"{'='*60}")

                extracted = self._extract_sheet_data(rows)
                output_parts.append(extracted)

            except Exception as e:
                output_parts.append(f"\n[Sheet '{sheet_name}' could not be parsed: {e}]")
                continue

        return "\n".join(output_parts)

    def _extract_sheet_data(self, rows: list) -> str:
        """
        Intelligently extract sheet data by detecting layout type:
        - Key-value pairs (label in col A, value in col B)
        - Tables (header row + data rows)
        - Mixed (sections with headers followed by key-value blocks)
        """
        parts = []
        i = 0

        while i < len(rows):
            row = rows[i]

            # Single cell row — treat as section header or standalone value
            if len(row) == 1:
                parts.append(f"\n[{row[0]}]")
                i += 1
                continue

            # Two-column row where col A looks like a label (ends with : or is short)
            # and col B has a value — key-value pair
            if len(row) >= 2 and row[0] and row[1]:
                # Check if the next several rows also look like key-value pairs
                kv_block = self._try_extract_kv_block(rows, i)
                if kv_block:
                    parts.append("")
                    for k, v in kv_block:
                        parts.append(f"  {k}: {v}")
                    i += len(kv_block)
                    continue

            # Multi-column row — check if this is a table header
            table, rows_consumed = self._try_extract_table(rows, i)
            if table:
                parts.append("")
                parts.append(table)
                i += rows_consumed
                continue

            # Fallback: join non-empty cells
            line = "  |  ".join(c for c in row if c)
            if line:
                parts.append(line)
            i += 1

        return "\n".join(parts)

    def _try_extract_kv_block(self, rows: list, start: int) -> list:
        """
        Check if rows starting at `start` form a key-value block.
        Returns list of (key, value) tuples if yes, empty list if no.
        """
        kv = []
        i = start
        while i < len(rows) and i < start + 50:
            row = rows[i]
            # Stop if row has more than 4 non-empty cells (looks like a table row)
            non_empty = [c for c in row if c]
            if len(non_empty) > 4:
                break
            if len(non_empty) == 0:
                i += 1
                continue
            if len(row) >= 2 and row[0] and row[1]:
                kv.append((row[0].rstrip(":"), row[1]))
            elif len(row) == 1 and row[0]:
                # Single label with no value — section sub-header, stop block
                break
            else:
                break
            i += 1
        # Only treat as a kv block if we found at least 2 pairs
        return kv if len(kv) >= 2 else []

    def _try_extract_table(self, rows: list, start: int) -> tuple:
        """
        Check if rows starting at `start` form a table (header + data rows).
        Returns formatted table string if yes, empty string if no.
        """
        if start >= len(rows):
            return "", 0
        header_row = rows[start]
        non_empty_headers = [c for c in header_row if c]
        if len(non_empty_headers) < 2:
            return "", 0

        # Collect data rows that have the same column count
        data_rows = []
        i = start + 1
        while i < len(rows) and i < start + 100:
            row = rows[i]
            non_empty = [c for c in row if c]
            if not non_empty:
                i += 1
                continue
            # Stop if this row looks like a new section header (single bold-like cell)
            if len(non_empty) == 1 and len(row) == 1:
                break
            data_rows.append(row)
            i += 1

        rows_consumed = (i - start)  # rows scanned past header
        if not data_rows:
            return "", 0

        # Format as readable table
        # Pad columns to header count
        col_count = len(header_row)
        lines = []
        lines.append("  " + " | ".join(str(h) for h in header_row if h))
        lines.append("  " + "-" * 60)
        for dr in data_rows:
            padded = list(dr) + [""] * (col_count - len(dr))
            values = [str(c) if c else "" for c in padded[:col_count]]
            line = " | ".join(v for v in values if v or True)
            # Only include rows with at least one value
            if any(v.strip() for v in values):
                lines.append("  " + line)
        return "\n".join(lines), rows_consumed
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        doc = fitz.open(file_path)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_parts.append(f"=== Page {page_num + 1} ===\n")
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        return "\n".join(text_parts)
