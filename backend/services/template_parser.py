"""
Template Parser Service

Parses Word (.docx) templates to extract:
- Section hierarchy (headings)
- Placeholders ({{variable_name}} format)
- Instructions for Claude (marked with special tags)
- Formatting metadata
"""

from docx import Document
from docx.shared import Pt
from typing import Dict, List, Any, Optional
import re
import json


class TemplateParser:
    """Parse Word document templates to extract structure"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        
    def parse(self) -> Dict[str, Any]:
        """
        Parse the template and return structured data
        
        Returns:
            Dict with:
            - sections: List of sections with headings
            - placeholders: List of all placeholders found
            - instructions: List of Claude instructions
            - metadata: Document metadata
        """
        return {
            "sections": self._extract_sections(),
            "placeholders": self._extract_placeholders(),
            "instructions": self._extract_instructions(),
            "metadata": self._extract_metadata()
        }
    
    def _extract_sections(self) -> List[Dict[str, Any]]:
        """
        Extract document sections based on heading styles
        
        Returns list of sections with:
        - level: Heading level (1-9)
        - title: Section title text
        - content: Text content under this heading
        - placeholders: Placeholders in this section
        """
        sections = []
        current_section = None
        
        for para in self.document.paragraphs:
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_section:
                    sections.append(current_section)
                
                # Extract heading level (e.g., "Heading 1" -> 1)
                level_match = re.search(r'Heading (\d+)', para.style.name)
                level = int(level_match.group(1)) if level_match else 1
                
                # Start new section
                current_section = {
                    "level": level,
                    "title": para.text.strip(),
                    "content": [],
                    "placeholders": []
                }
            elif current_section is not None:
                # Add content to current section
                text = para.text.strip()
                if text:
                    current_section["content"].append(text)
                    # Extract placeholders from this paragraph
                    placeholders = self._find_placeholders(text)
                    current_section["placeholders"].extend(placeholders)
        
        # Add last section
        if current_section:
            sections.append(current_section)
        
        # Convert content lists to strings
        for section in sections:
            section["content"] = "\n".join(section["content"])
        
        return sections
    
    def _extract_placeholders(self) -> List[Dict[str, Any]]:
        """
        Extract all placeholders in {{variable_name}} format
        
        Returns list of unique placeholders with:
        - name: Variable name
        - occurrences: Number of times it appears
        - sections: Which sections contain it
        """
        placeholder_map = {}
        
        for para in self.document.paragraphs:
            placeholders = self._find_placeholders(para.text)
            for ph in placeholders:
                if ph not in placeholder_map:
                    placeholder_map[ph] = {
                        "name": ph,
                        "occurrences": 0,
                        "sections": []
                    }
                placeholder_map[ph]["occurrences"] += 1
        
        return list(placeholder_map.values())
    
    def _find_placeholders(self, text: str) -> List[str]:
        """Find all {{placeholder}} patterns in text"""
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(pattern, text)
        return [match.strip() for match in matches]
    
    def _extract_instructions(self) -> List[Dict[str, Any]]:
        """
        Extract Claude instructions marked with special tags
        
        Looks for patterns like:
        [CLAUDE: instruction text]
        or
        <!-- CLAUDE: instruction text -->
        
        Returns list of instructions with:
        - text: Instruction content
        - section: Which section it's in
        - type: Type of instruction (tone, style, format, etc.)
        """
        instructions = []
        current_section = "Unknown"
        
        for para in self.document.paragraphs:
            # Track current section
            if para.style.name.startswith('Heading'):
                current_section = para.text.strip()
            
            # Look for instruction patterns
            text = para.text
            
            # Pattern 1: [CLAUDE: ...]
            claude_pattern = r'\[CLAUDE:\s*([^\]]+)\]'
            matches = re.findall(claude_pattern, text, re.IGNORECASE)
            for match in matches:
                instructions.append({
                    "text": match.strip(),
                    "section": current_section,
                    "type": self._classify_instruction(match)
                })
            
            # Pattern 2: <!-- CLAUDE: ... -->
            comment_pattern = r'<!--\s*CLAUDE:\s*([^-]+)-->'
            matches = re.findall(comment_pattern, text, re.IGNORECASE)
            for match in matches:
                instructions.append({
                    "text": match.strip(),
                    "section": current_section,
                    "type": self._classify_instruction(match)
                })
        
        return instructions
    
    def _classify_instruction(self, text: str) -> str:
        """Classify instruction type based on keywords"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['tone', 'voice', 'style']):
            return 'tone'
        elif any(word in text_lower for word in ['format', 'structure', 'layout']):
            return 'format'
        elif any(word in text_lower for word in ['length', 'words', 'sentences']):
            return 'length'
        elif any(word in text_lower for word in ['technical', 'detail', 'depth']):
            return 'detail'
        else:
            return 'general'
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """
        Extract document metadata
        
        Returns:
        - total_paragraphs: Number of paragraphs
        - total_sections: Number of heading sections
        - total_placeholders: Total placeholder count
        - styles_used: List of styles used in document
        """
        styles_used = set()
        for para in self.document.paragraphs:
            styles_used.add(para.style.name)
        
        return {
            "total_paragraphs": len(self.document.paragraphs),
            "total_sections": len([p for p in self.document.paragraphs if p.style.name.startswith('Heading')]),
            "total_placeholders": len(self._extract_placeholders()),
            "styles_used": sorted(list(styles_used))
        }


def parse_template_file(file_path: str) -> Dict[str, Any]:
    """
    Convenience function to parse a template file
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Parsed template structure
    """
    parser = TemplateParser(file_path)
    return parser.parse()
